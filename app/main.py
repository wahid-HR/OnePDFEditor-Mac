#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
One PDF Editor v1.0
Made by Yumdrop Tech. Studio – 2026
Offline PDF / Image / Document viewer & editor for Windows.
"""

import os
import queue
import sys
import io
import tempfile
import webbrowser
from pathlib import Path
from typing import Optional, List, Tuple, Dict, Any
from datetime import datetime

try:
    import pymupdf as fitz
except ImportError:
    import fitz  # type: ignore

from PIL import Image, ImageDraw, ImageTk, ImageEnhance, ImageFilter
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog, colorchooser

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def convert_docx_via_word(path):
    """Convert .docx to PDF using installed Microsoft Word (best fidelity).
    Returns path to a temp PDF. Raises on failure.
    """
    path = str(Path(path).resolve())
    if not path.lower().endswith(".docx"):
        raise RuntimeError("Not a Word file")
    tmp_pdf = str(Path(tempfile.gettempdir()) / f"onepdf_docx_{os.getpid()}_{abs(hash(path)) % 10**8}.pdf")
    if os.path.exists(tmp_pdf):
        try:
            os.remove(tmp_pdf)
        except Exception:
            pass
    word = None
    doc = None
    try:
        import win32com.client  # type: ignore
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(path, ReadOnly=True)
        # 17 = wdFormatPDF
        doc.SaveAs(tmp_pdf, FileFormat=17)
        doc.Close(False)
        doc = None
        word.Quit()
        word = None
        if not os.path.isfile(tmp_pdf):
            raise RuntimeError("Word did not produce a PDF")
        return tmp_pdf
    except Exception as e:
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        raise RuntimeError(
            "High-quality Word→PDF needs Microsoft Word installed.\\n"
            f"Details: {e}"
        )


def convert_docx_simple(path):
    """Fallback: paragraphs only (no tables/images/headers)."""
    if not HAS_DOCX:
        raise RuntimeError("python-docx not available")
    document = DocxDocument(path)
    pdf = fitz.open()
    page = pdf.new_page(width=595, height=842)
    y = 50
    for para in document.paragraphs:
        text = (para.text or "").strip()
        if not text:
            y += 8
            continue
        while text:
            chunk = text[:90]
            text = text[90:]
            if y > 800:
                page = pdf.new_page(width=595, height=842)
                y = 50
            page.insert_text((50, y), chunk, fontsize=11, fontname="helv")
            y += 16
    return pdf


def docx_to_fitz(path):
    """Prefer MS Word COM (keeps tables/images/headers); else simple text PDF."""
    if sys.platform.startswith("win"):
        try:
            pdf_path = convert_docx_via_word(path)
            doc = fitz.open(pdf_path)
            # load into memory so temp can be deleted
            data = doc.tobytes()
            page_count = doc.page_count
            doc.close()
            try:
                os.remove(pdf_path)
            except Exception:
                pass
            return fitz.open("pdf", data)
        except Exception:
            pass
    return convert_docx_simple(path)


APP_NAME = "One PDF Editor"
APP_VERSION = "1.0"
APP_YEAR = "2026"
APP_STUDIO = "Yumdrop Tech. Studio"


def resource_path(*parts):
    """Resolve path for source, onedir, and onefile PyInstaller builds."""
    candidates = []
    if getattr(sys, "frozen", False):
        # onefile / onedir internal
        if hasattr(sys, "_MEIPASS"):
            candidates.append(Path(sys._MEIPASS))
        # onedir: next to the .exe
        candidates.append(Path(sys.executable).resolve().parent)
        # sometimes assets sit under _internal next to exe
        candidates.append(Path(sys.executable).resolve().parent / "_internal")
    else:
        candidates.append(Path(__file__).resolve().parent.parent)
        candidates.append(Path(__file__).resolve().parent)
    for base in candidates:
        p = base.joinpath(*parts)
        if p.exists():
            return p
    # fallback last candidate
    return candidates[0].joinpath(*parts)

MAX_UNDO = 12
MIN_FONT_SIZE = 4.0
DEFAULT_ZOOM = 1.0

BASE_FONTS = {
    "helv": "helv", "helvetica": "helv", "arial": "helv", "sans": "helv",
    "times": "times", "times-roman": "times", "serif": "times",
    "cour": "cour", "courier": "cour", "mono": "cour",
}

COLORS = {
    "bg": "#1e1e2e",
    "surface": "#2a2a3c",
    "surface2": "#34344a",
    "accent": "#7c5cff",
    "accent_hover": "#9b82ff",
    "text": "#e8e8f0",
    "text_dim": "#a0a0b8",
    "canvas_bg": "#12121a",
    "toolbar": "#252536",
}


def safe_color_to_rgb(color_int):
    if color_int is None:
        return (0.0, 0.0, 0.0)
    r = ((color_int >> 16) & 0xFF) / 255.0
    g = ((color_int >> 8) & 0xFF) / 255.0
    b = (color_int & 0xFF) / 255.0
    return (r, g, b)


def flags_to_style(flags):
    return {
        "italic": bool(flags & 2**1),
        "bold": bool(flags & 2**4),
        "mono": bool(flags & 2**3),
        "serif": bool(flags & 2**2),
    }


def choose_fallback_font(original_name, flags, text=""):
    name = (original_name or "").lower()
    style = flags_to_style(flags)
    if "+" in name:
        name = name.split("+", 1)[-1]
    for key, alias in BASE_FONTS.items():
        if key in name:
            return alias
    if style["mono"]:
        return "cour"
    if style["serif"] or "times" in name or "roman" in name:
        return "times"
    return "helv"


def text_needs_complex_script(text):
    for ch in text:
        o = ord(ch)
        if o > 0x00FF and not (0x2000 <= o <= 0x206F) and not (0x2010 <= o <= 0x2027):
            return True
    return False


def text_has_bengali(text):
    """True if text contains Bengali Unicode letters."""
    for ch in text:
        o = ord(ch)
        if 0x0980 <= o <= 0x09FF:  # Bengali block
            return True
        if 0x200C <= o <= 0x200D:  # ZWNJ/ZWJ often used in Bangla
            continue
    return False


# Minimal Bijoy (ANSI) → Unicode map for common Bangla typing
# Users typing with Bijoy keyboard layout on Windows often produce these codes.
_BIJOY_MAP = {
    "A": "অ", "i": "ই", "I": "ঈ", "u": "উ", "U": "ঊ", "e": "এ", "E": "ঐ",
    "o": "ও", "O": "ঔ", "k": "ক", "K": "খ", "g": "গ", "G": "ঘ", "c": "চ",
    "C": "ছ", "j": "জ", "J": "ঝ", "T": "ট", "t": "ত", "d": "দ", "D": "ড",
    "n": "ন", "N": "ণ", "p": "প", "P": "ফ", "f": "ফ", "b": "ব", "v": "ভ",
    "m": "ম", "z": "য", "r": "র", "l": "ল", "S": "শ", "s": "স", "h": "হ",
    "R": "ড়", "y": "য়", "w": "ৎ", "Y": "য়",
    "a": "া", "w": "্", "x": "ঁ", "X": "ঁ",
    "1": "১", "2": "২", "3": "৩", "4": "৪", "5": "৫",
    "6": "৬", "7": "৭", "8": "৮", "9": "৯", "0": "০",
}


def bijoy_to_unicode(text, force=False):
    """Best-effort Bijoy→Unicode. Only when force=True (user chose Bijoy mode).
    Never auto-convert normal English text.
    """
    if text_has_bengali(text):
        return text
    if not force:
        return text
    out = []
    for ch in text:
        out.append(_BIJOY_MAP.get(ch, ch))
    return "".join(out)


def get_bengali_font_path(bold=False):
    name = "NotoSansBengali-Bold.ttf" if bold else "NotoSansBengali-Regular.ttf"
    p = resource_path("assets", "fonts", name)
    if p.exists():
        return str(p)
    # try sibling
    p2 = resource_path("fonts", name)
    if p2.exists():
        return str(p2)
    return None


def measure_text_width(text, fontname, fontsize):
    try:
        font = fitz.Font(fontname)
        return font.text_length(text, fontsize=fontsize)
    except Exception:
        return len(text) * fontsize * 0.5


class PDFDocument:
    def __init__(self):
        self.doc = None
        self.path = None
        self.dirty = False
        self.source_type = "pdf"
        self._undo_stack = []
        self._redo_stack = []

    def open(self, path, password=""):
        ext = Path(path).suffix.lower()
        self.close()
        try:
            if ext in (".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp", ".tiff", ".tif"):
                return self._open_image(path)
            if ext == ".docx" and HAS_DOCX:
                return self._open_docx(path)
            self.doc = fitz.open(path)
            if self.doc.is_encrypted:
                if not self.doc.authenticate(password):
                    self.doc.close()
                    self.doc = None
                    return False
            self.path = path
            self.source_type = "pdf"
            self.dirty = False
            self._undo_stack.clear()
            self._push_undo()
            return True
        except Exception as e:
            raise RuntimeError(f"Cannot open file: {e}") from e

    def _open_image(self, path):
        img = Image.open(path)
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        # Offline quality enhance
        try:
            img = ImageEnhance.Sharpness(img).enhance(1.35)
            img = ImageEnhance.Contrast(img).enhance(1.12)
            img = ImageEnhance.Color(img).enhance(1.05)
        except Exception:
            pass
        bio = io.BytesIO()
        img.save(bio, format="PDF", resolution=150.0)
        bio.seek(0)
        self.doc = fitz.open("pdf", bio.read())
        # Best-effort OCR text layer (needs Tesseract on system)
        try:
            page = self.doc[0]
            tp = page.get_textpage_ocr(dpi=150, full=True)
            # OCR text becomes selectable/editable spans
            _ = page.get_text("dict", textpage=tp)
        except Exception:
            pass  # no tesseract — image-only PDF still works
        self.path = path
        self.source_type = "image"
        self.dirty = False
        self._undo_stack.clear()
        self._push_undo()
        return True

    def _open_docx(self, path):
        self.doc = docx_to_fitz(path)
        self.path = path
        self.source_type = "docx"
        self.dirty = True
        self._undo_stack.clear()
        self._push_undo()
        return True

    def close(self):
        if self.doc:
            try:
                self.doc.close()
            except Exception:
                pass
        self.doc = None
        self.path = None
        self.dirty = False
        self.source_type = "pdf"
        self._undo_stack.clear()
        self._redo_stack.clear()

    def page_count(self):
        return len(self.doc) if self.doc else 0

    def get_page_links(self, page_idx):
        page = self.get_page(page_idx)
        if not page:
            return []
        try:
            return list(page.get_links())
        except Exception:
            return []

    def update_link_uri(self, page_idx, link, new_uri):
        page = self.get_page(page_idx)
        if not page or not link:
            return False, "Invalid"
        self.save_state()
        try:
            new_uri = (new_uri or "").strip()
            if not new_uri:
                return False, "Empty URL"
            if not (new_uri.startswith("http://") or new_uri.startswith("https://")
                    or new_uri.startswith("mailto:") or new_uri.startswith("ftp://")):
                if "@" in new_uri and " " not in new_uri:
                    new_uri = "mailto:" + new_uri
                else:
                    new_uri = "https://" + new_uri
            # Match by xref or rect against live links
            live = page.get_links()
            target = None
            want_xref = link.get("xref")
            want_from = fitz.Rect(link.get("from"))
            for lk in live:
                if want_xref and lk.get("xref") == want_xref:
                    target = lk
                    break
                if fitz.Rect(lk.get("from")).irect == want_from.irect:
                    target = lk
                    break
            if target is None and live:
                # fallback: first URI link with same uri
                for lk in live:
                    if lk.get("uri") == link.get("uri"):
                        target = lk
                        break
            if target is None:
                return False, "Link not found on page (try again)"
            rect = fitz.Rect(target.get("from"))
            page.delete_link(target)
            page.insert_link({
                "kind": fitz.LINK_URI,
                "from": rect,
                "uri": new_uri,
            })
            self.dirty = True
            return True, "Link updated"
        except Exception as e:
            return False, str(e)

    def delete_link(self, page_idx, link):
        page = self.get_page(page_idx)
        if not page or not link:
            return False, "Invalid"
        self.save_state()
        try:
            page.delete_link(dict(link))
            self.dirty = True
            return True, "Link removed"
        except Exception as e:
            return False, str(e)

    def add_uri_link(self, page_idx, rect, uri):
        page = self.get_page(page_idx)
        if not page:
            return False, "Invalid page"
        uri = (uri or "").strip()
        if not uri:
            return False, "Empty URL"
        if not (uri.startswith("http://") or uri.startswith("https://") or uri.startswith("mailto:")):
            if "@" in uri and " " not in uri:
                uri = "mailto:" + uri
            else:
                uri = "https://" + uri
        self.save_state()
        try:
            r = fitz.Rect(rect)
            r.normalize()
            if r.width < 3 or r.height < 3:
                return False, "Area too small"
            page.insert_link({
                "kind": fitz.LINK_URI,
                "from": r,
                "uri": uri,
            })
            self.dirty = True
            return True, "Link added"
        except Exception as e:
            return False, str(e)


    def get_page(self, idx):
        if self.doc and 0 <= idx < len(self.doc):
            return self.doc[idx]
        return None

    def _serialize(self):
        if not self.doc:
            return b""
        return self.doc.tobytes(garbage=3, deflate=True)

    def _push_undo(self):
        data = self._serialize()
        if data:
            self._undo_stack.append(data)
            if len(self._undo_stack) > MAX_UNDO:
                self._undo_stack.pop(0)
            self._redo_stack.clear()

    def save_state(self):
        self._push_undo()
        self.dirty = True

    def undo(self):
        if len(self._undo_stack) < 2 or not self.doc:
            return False
        current = self._undo_stack.pop()
        self._redo_stack.append(current)
        self._restore(self._undo_stack[-1])
        self.dirty = True
        return True

    def redo(self):
        if not self._redo_stack or not self.doc:
            return False
        data = self._redo_stack.pop()
        self._undo_stack.append(data)
        self._restore(data)
        self.dirty = True
        return True

    def _restore(self, data):
        if not data:
            return
        try:
            newdoc = fitz.open("pdf", data)
            if self.doc:
                self.doc.close()
            self.doc = newdoc
        except Exception:
            pass

    def save(self, path=None):
        if not self.doc:
            return False
        target = path or self.path
        if not target:
            return False
        if not target.lower().endswith(".pdf"):
            target = str(Path(target).with_suffix(".pdf"))
        try:
            self.doc.save(target, garbage=3, deflate=True, incremental=False)
            self.path = target
            self.dirty = False
            self.source_type = "pdf"
            return True
        except Exception as e:
            raise RuntimeError(f"Save failed: {e}") from e

    def get_text_spans(self, page_idx):
        page = self.get_page(page_idx)
        if not page:
            return []
        spans = []
        try:
            blocks = page.get_text("dict", flags=fitz.TEXTFLAGS_TEXT)["blocks"]
            for b in blocks:
                if b.get("type") != 0:
                    continue
                for line in b.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip("\x00")
                        if not text.strip():
                            continue
                        spans.append({
                            "text": text,
                            "bbox": fitz.Rect(span["bbox"]),
                            "origin": fitz.Point(span.get("origin", span["bbox"][:2])),
                            "font": span.get("font", "helv"),
                            "size": float(span.get("size", 11)),
                            "color": span.get("color", 0),
                            "flags": span.get("flags", 0),
                        })
        except Exception:
            pass
        return spans

    def search(self, query):
        results = []
        if not self.doc or not query:
            return results
        for i in range(len(self.doc)):
            try:
                for r in self.doc[i].search_for(query, quads=False):
                    results.append((i, fitz.Rect(r)))
            except Exception:
                continue
        return results

    def replace_span(self, page_idx, span, new_text, fontsize=None, bold=False, italic=False):
        page = self.get_page(page_idx)
        if not page or not new_text:
            return False, "Invalid"
        self.save_state()
        bbox = fitz.Rect(span["bbox"])
        origin = fitz.Point(span["origin"])
        if fontsize is None:
            fontsize = float(span["size"])
        color = safe_color_to_rgb(span["color"])
        flags = span.get("flags", 0)
        if bold:
            flags |= 2**4
        if italic:
            flags |= 2**1
        fontname = span.get("font", "helv")
        use_font = choose_fallback_font(fontname, flags, new_text)
        # Prefer styled base fonts when bold/italic requested
        if bold and italic:
            if use_font == "helv":
                use_font = "heit"  # Helvetica-Oblique approx; fall back
            # Base14 has limited styled names; use morph or just size
        complex_script = text_needs_complex_script(new_text)
        try:
            _ = fitz.Font(use_font)
        except Exception:
            use_font = "helv"
        # Honor user-requested size (do not auto-shrink when size is set from UI)
        fitted_size = max(MIN_FONT_SIZE, min(96.0, float(fontsize)))
        # Sample approximate background from page render near the span
        bg_fill = (1, 1, 1)
        try:
            pix = page.get_pixmap(clip=bbox, matrix=fitz.Matrix(2, 2), alpha=False)
            if pix.width > 2 and pix.height > 2 and pix.samples:
                # average a few corner pixels (avoid center which is text)
                samples = []
                w, h = pix.width, pix.height
                for (px, py) in [(1, 1), (w-2, 1), (1, h-2), (w-2, h-2), (w//2, 1), (w//2, h-2)]:
                    i = (py * w + px) * 3
                    if i + 2 < len(pix.samples):
                        samples.append((pix.samples[i]/255.0, pix.samples[i+1]/255.0, pix.samples[i+2]/255.0))
                if samples:
                    bg_fill = (
                        sum(s[0] for s in samples) / len(samples),
                        sum(s[1] for s in samples) / len(samples),
                        sum(s[2] for s in samples) / len(samples),
                    )
        except Exception:
            pass
        try:
            page.add_redact_annot(bbox, fill=bg_fill)
            page.apply_redactions(
                images=fitz.PDF_REDACT_IMAGE_NONE,
                graphics=fitz.PDF_REDACT_LINE_ART_NONE,
                text=fitz.PDF_REDACT_TEXT_REMOVE,
            )
        except Exception as e:
            return False, f"Redaction failed: {e}"
        # Do NOT auto-convert English to Bangla. Use Bangla font only if text has Bangla letters.
        is_bn = text_has_bengali(new_text)
        fontfile = get_bengali_font_path(bold=bold) if is_bn else None

        try:
            if fontfile:
                # Register & use embedded Unicode font (Bangla etc.)
                try:
                    page.insert_font(fontname="bnfont", fontfile=fontfile)
                    fname = "bnfont"
                except Exception:
                    fname = use_font
                page.insert_text(
                    origin, new_text, fontname=fname, fontfile=fontfile,
                    fontsize=fitted_size, color=color, render_mode=0, overlay=True,
                )
                if bold:
                    page.insert_text(
                        fitz.Point(origin.x + 0.4, origin.y), new_text,
                        fontname=fname, fontfile=fontfile,
                        fontsize=fitted_size, color=color, render_mode=0, overlay=True,
                    )
            else:
                page.insert_text(
                    origin, new_text, fontname=use_font, fontsize=fitted_size,
                    color=color, render_mode=0, overlay=True,
                )
                if bold:
                    page.insert_text(
                        fitz.Point(origin.x + 0.35, origin.y), new_text,
                        fontname=use_font, fontsize=fitted_size,
                        color=color, render_mode=0, overlay=True,
                    )
        except Exception as e:
            # Last resort: textbox with fontfile
            try:
                if fontfile:
                    page.insert_textbox(
                        bbox, new_text, fontname="helv", fontfile=fontfile,
                        fontsize=fitted_size, color=color, align=0, overlay=True,
                    )
                else:
                    return False, f"Insert failed: {e}"
            except Exception as e2:
                return False, f"Insert failed: {e2}"
        self.dirty = True
        msg = f"Edited (size={fitted_size:.1f}"
        if bold:
            msg += ", bold"
        if is_bn:
            msg += ", Bangla"
        msg += ")"
        return True, msg


    def extract_page_images(self, page_idx):
        """Return list of (xref, PIL.Image) for images on the page."""
        page = self.get_page(page_idx)
        if not page or not self.doc:
            return []
        result = []
        try:
            for info in page.get_images(full=True):
                xref = info[0]
                try:
                    pix = fitz.Pixmap(self.doc, xref)
                    if pix.n - pix.alpha > 3:  # CMYK etc
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    mode = "RGBA" if pix.alpha else "RGB"
                    img = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
                    result.append((xref, img))
                except Exception:
                    continue
        except Exception:
            pass
        return result

    def sample_rect_background(self, page_idx, rect):
        """Average background color from corners of a page region (0-1 RGB)."""
        page = self.get_page(page_idx)
        if not page:
            return (1.0, 1.0, 1.0)
        try:
            # slightly inset to avoid borders
            r = fitz.Rect(rect)
            if r.width < 2 or r.height < 2:
                return (1.0, 1.0, 1.0)
            pix = page.get_pixmap(clip=r, matrix=fitz.Matrix(2, 2), alpha=False)
            w, h = pix.width, pix.height
            if w < 2 or h < 2 or not pix.samples:
                return (1.0, 1.0, 1.0)
            pts = [(1, 1), (w - 2, 1), (1, h - 2), (w - 2, h - 2),
                   (w // 2, 1), (w // 2, h - 2), (1, h // 2), (w - 2, h // 2)]
            samples = []
            for px, py in pts:
                i = (py * w + px) * 3
                if i + 2 < len(pix.samples):
                    samples.append((
                        pix.samples[i] / 255.0,
                        pix.samples[i + 1] / 255.0,
                        pix.samples[i + 2] / 255.0,
                    ))
            if not samples:
                return (1.0, 1.0, 1.0)
            return (
                sum(s[0] for s in samples) / len(samples),
                sum(s[1] for s in samples) / len(samples),
                sum(s[2] for s in samples) / len(samples),
            )
        except Exception:
            return (1.0, 1.0, 1.0)

    def fill_rectangle(self, page_idx, rect, color_rgb=None):
        """Cover area to hide text. If color_rgb is None, auto-sample PDF background."""
        page = self.get_page(page_idx)
        if not page:
            return False
        if color_rgb is None:
            color_rgb = self.sample_rect_background(page_idx, rect)
        self.save_state()
        try:
            page.add_redact_annot(rect, fill=color_rgb)
            page.apply_redactions(
                images=fitz.PDF_REDACT_IMAGE_NONE,
                graphics=fitz.PDF_REDACT_LINE_ART_NONE,
                text=fitz.PDF_REDACT_TEXT_REMOVE,
            )
            shape = page.new_shape()
            shape.draw_rect(rect)
            shape.finish(color=color_rgb, fill=color_rgb, width=0)
            shape.commit()
            self.dirty = True
            return True
        except Exception:
            return False

    def move_span(self, page_idx, span, dx, dy):
        """Move text span by (dx, dy) in page coordinates. Keeps text/size/color."""
        page = self.get_page(page_idx)
        if not page or not span:
            return False, "Invalid"
        if abs(dx) < 0.5 and abs(dy) < 0.5:
            return False, "No move"
        self.save_state()
        bbox = fitz.Rect(span["bbox"])
        origin = fitz.Point(span["origin"])
        text = span.get("text") or ""
        if not text.strip():
            return False, "Empty"
        fontsize = float(span.get("size") or 11)
        color = safe_color_to_rgb(span.get("color", 0))
        flags = span.get("flags", 0)
        bold = bool(flags & 16)
        is_bn = text_has_bengali(text)
        fontfile = get_bengali_font_path(bold=bold) if is_bn else None
        fontname = choose_fallback_font(span.get("font", "helv"), flags, text)
        try:
            _ = fitz.Font(fontname)
        except Exception:
            fontname = "helv"
        # Sample background from corners/edges (avoid center text ink)
        bg_fill = self.sample_rect_background(page_idx, bbox)
        try:
            # Slightly pad redact so no leftover glyph edges
            pad = fitz.Rect(bbox.x0 - 0.5, bbox.y0 - 0.5, bbox.x1 + 0.5, bbox.y1 + 0.5)
            page.add_redact_annot(pad, fill=bg_fill)
            page.apply_redactions(
                images=fitz.PDF_REDACT_IMAGE_NONE,
                graphics=fitz.PDF_REDACT_LINE_ART_NONE,
                text=fitz.PDF_REDACT_TEXT_REMOVE,
            )
            # Paint solid bg rect so hole matches page
            shape = page.new_shape()
            shape.draw_rect(pad)
            shape.finish(color=bg_fill, fill=bg_fill, width=0)
            shape.commit()
            new_origin = fitz.Point(origin.x + dx, origin.y + dy)
            if fontfile:
                try:
                    page.insert_font(fontname="bnfont", fontfile=fontfile)
                    page.insert_text(new_origin, text, fontname="bnfont", fontfile=fontfile,
                                     fontsize=fontsize, color=color, overlay=True)
                except Exception:
                    page.insert_text(new_origin, text, fontname=fontname, fontsize=fontsize,
                                     color=color, overlay=True)
            else:
                page.insert_text(new_origin, text, fontname=fontname, fontsize=fontsize,
                                 color=color, overlay=True)
            self.dirty = True
            return True, "Moved"
        except Exception as e:
            return False, str(e)

    def insert_text_in_rect(self, page_idx, rect, text, fontsize=None, color=(0, 0, 0)):
        page = self.get_page(page_idx)
        if not page or not text:
            return False
        self.save_state()
        try:
            r = fitz.Rect(rect)
            if fontsize is None:
                fontsize = max(6.0, min(24.0, r.height * 0.65))
            origin = fitz.Point(r.x0 + 1, r.y0 + fontsize * 0.85)
            fontfile = get_bengali_font_path() if text_has_bengali(text) else None
            kwargs = dict(fontsize=fontsize, color=color, render_mode=0, overlay=True)
            if fontfile:
                try:
                    page.insert_font(fontname="bnfont", fontfile=fontfile)
                except Exception:
                    pass
                page.insert_text(origin, text, fontname="bnfont", fontfile=fontfile, **kwargs)
            else:
                page.insert_text(origin, text, fontname="helv", **kwargs)
            self.dirty = True
            return True
        except Exception:
            return False

    def insert_signature_image(self, page_idx, img, rect):
        page = self.get_page(page_idx)
        if not page:
            return False
        self.save_state()
        try:
            if img.mode != "RGBA":
                img = img.convert("RGBA")
            bio = io.BytesIO()
            img.save(bio, format="PNG")
            bio.seek(0)
            page.insert_image(rect, stream=bio.getvalue(), overlay=True)
            self.dirty = True
            return True
        except Exception:
            return False


class SignatureDialog(tk.Toplevel):
    def __init__(self, parent, on_done):
        super().__init__(parent)
        self.title("Draw Signature")
        self.configure(bg=COLORS["surface"])
        self.resizable(True, True)
        self.transient(parent)
        self.grab_set()
        self.on_done = on_done
        self.strokes = []
        self.current_stroke = []
        self.drawing = False
        self.canvas = tk.Canvas(self, bg="white", width=520, height=220, cursor="pencil",
                                highlightthickness=1, highlightbackground=COLORS["accent"])
        self.canvas.pack(fill=tk.BOTH, expand=True, padx=12, pady=12)
        self.canvas.bind("<ButtonPress-1>", self._start)
        self.canvas.bind("<B1-Motion>", self._move)
        self.canvas.bind("<ButtonRelease-1>", self._end)
        btn_frame = tk.Frame(self, bg=COLORS["surface"])
        btn_frame.pack(fill=tk.X, padx=12, pady=(0, 12))
        self._btn(btn_frame, "Clear", self._clear).pack(side=tk.LEFT, padx=4)
        self._btn(btn_frame, "Cancel", self._cancel).pack(side=tk.RIGHT, padx=4)
        self._btn(btn_frame, "Use Signature", self._accept, accent=True).pack(side=tk.RIGHT, padx=4)
        self.protocol("WM_DELETE_WINDOW", self._cancel)
        self.geometry("560x320")
        self.focus_set()

    def _btn(self, parent, text, cmd, accent=False):
        bg = COLORS["accent"] if accent else COLORS["surface2"]
        return tk.Button(parent, text=text, command=cmd, bg=bg, fg=COLORS["text"],
                         activebackground=COLORS["accent_hover"], activeforeground="white",
                         relief=tk.FLAT, padx=14, pady=6, cursor="hand2", font=("Segoe UI", 10))

    def _start(self, event):
        self.drawing = True
        self.current_stroke = [(event.x, event.y)]
        self.canvas.create_oval(event.x-1, event.y-1, event.x+1, event.y+1, fill="black", outline="black")

    def _move(self, event):
        if not self.drawing:
            return
        x, y = event.x, event.y
        if self.current_stroke:
            px, py = self.current_stroke[-1]
            self.canvas.create_line(px, py, x, y, fill="black", width=2, capstyle=tk.ROUND, smooth=True)
        self.current_stroke.append((x, y))

    def _end(self, event):
        if self.drawing and self.current_stroke:
            self.strokes.append(self.current_stroke)
        self.drawing = False
        self.current_stroke = []

    def _clear(self):
        self.canvas.delete("all")
        self.strokes.clear()
        self.current_stroke = []

    def _cancel(self):
        self.on_done(None)
        self.destroy()

    def _accept(self):
        if not self.strokes:
            messagebox.showwarning("Empty", "Please draw a signature first.", parent=self)
            return
        w = int(self.canvas.winfo_width()) or 520
        h = int(self.canvas.winfo_height()) or 220
        img = Image.new("RGBA", (w, h), (0, 0, 0, 0))
        draw = ImageDraw.Draw(img)
        for stroke in self.strokes:
            if len(stroke) < 2:
                if stroke:
                    x, y = stroke[0]
                    draw.ellipse([x-1, y-1, x+1, y+1], fill=(0, 0, 0, 255))
                continue
            draw.line(stroke, fill=(0, 0, 0, 255), width=2, joint="curve")
        self.on_done(img)
        self.destroy()


class ScreenshotPopup(tk.Toplevel):
    def __init__(self, parent, image):
        super().__init__(parent)
        self.title("Screenshot")
        self.configure(bg=COLORS["surface"])
        self.transient(parent)
        self.grab_set()
        self.image = image
        preview = image.copy()
        preview.thumbnail((480, 360))
        self.photo = ImageTk.PhotoImage(preview)
        lbl = tk.Label(self, image=self.photo, bg=COLORS["surface"])
        lbl.pack(padx=16, pady=16)
        btn_frame = tk.Frame(self, bg=COLORS["surface"])
        btn_frame.pack(pady=(0, 16))
        tk.Button(btn_frame, text="Copy to Clipboard", command=self._copy,
                  bg=COLORS["accent"], fg="white", relief=tk.FLAT, padx=16, pady=8,
                  font=("Segoe UI", 10), cursor="hand2").pack(side=tk.LEFT, padx=8)
        tk.Button(btn_frame, text="Save as PNG", command=self._save,
                  bg=COLORS["surface2"], fg=COLORS["text"], relief=tk.FLAT, padx=16, pady=8,
                  font=("Segoe UI", 10), cursor="hand2").pack(side=tk.LEFT, padx=8)
        tk.Button(btn_frame, text="Close", command=self.destroy,
                  bg=COLORS["surface2"], fg=COLORS["text_dim"], relief=tk.FLAT, padx=12, pady=8,
                  font=("Segoe UI", 10), cursor="hand2").pack(side=tk.LEFT, padx=8)
        self.geometry("520x460")
        self.focus_set()

    def _copy(self):
        try:
            output = io.BytesIO()
            self.image.convert("RGB").save(output, "BMP")
            data = output.getvalue()[14:]
            output.close()
            try:
                import win32clipboard
                win32clipboard.OpenClipboard()
                win32clipboard.EmptyClipboard()
                win32clipboard.SetClipboardData(win32clipboard.CF_DIB, data)
                win32clipboard.CloseClipboard()
                messagebox.showinfo("Copied", "Screenshot copied to clipboard.", parent=self)
            except ImportError:
                tmp = Path(tempfile.gettempdir()) / "one_pdf_screenshot.png"
                self.image.save(tmp)
                messagebox.showinfo("Saved", f"Clipboard helper not available.\nSaved to:\n{tmp}", parent=self)
        except Exception as e:
            messagebox.showerror("Error", str(e), parent=self)

    def _save(self):
        path = filedialog.asksaveasfilename(
            parent=self, defaultextension=".png",
            filetypes=[("PNG image", "*.png"), ("JPEG", "*.jpg")],
            initialfile=f"screenshot_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png",
        )
        if path:
            self.image.save(path)
            messagebox.showinfo("Saved", f"Saved to:\n{path}", parent=self)
            self.destroy()



class BanglaKeyboard(tk.Toplevel):
    """On-screen Bangla Unicode keyboard — no system IME required."""

    # Rows of Bangla characters for a compact keyboard
    ROWS = [
        list("১২৩৪৫৬৭৮৯০"),
        list("অআইঈউঊঋএঐওঔ"),
        list("কখগঘঙচছজঝঞ"),
        list("টঠডঢণতথদধন"),
        list("পফবভমযরল"),
        list("শষসহড়ঢ়য়ৎ"),
        list("ািীুূৃেৈোৌ্ংঃঁ"),
    ]
    EXTRA = [
        ("ক্ষ", "ক্ষ"), ("জ্ঞ", "জ্ঞ"), ("ঞ্জ", "ঞ্জ"), ("ত্ত", "ত্ত"),
        ("Space", " "), ("⌫", "BACK"), ("Enter", "ENTER"),
    ]

    def __init__(self, parent, on_char, on_enter=None):
        super().__init__(parent)
        self.title("বাংলা কিবোর্ড")
        self.configure(bg=COLORS["surface"])
        self.transient(parent)
        self.on_char = on_char
        self.on_enter = on_enter
        self.resizable(False, False)

        tk.Label(
            self, text="ক্লিক করে বাংলা লিখুন (Avro লাগবে না)",
            bg=COLORS["surface"], fg=COLORS["text_dim"], font=("Segoe UI", 9),
        ).pack(pady=(8, 4))

        # Phonetic quick type
        pf = tk.Frame(self, bg=COLORS["surface"])
        pf.pack(fill=tk.X, padx=10, pady=4)
        tk.Label(pf, text="Phonetic:", bg=COLORS["surface"], fg=COLORS["text"],
                 font=("Segoe UI", 9)).pack(side=tk.LEFT)
        self.phon_var = tk.StringVar()
        pe = tk.Entry(pf, textvariable=self.phon_var, width=22,
                      bg=COLORS["surface2"], fg=COLORS["text"],
                      insertbackground=COLORS["text"], relief=tk.FLAT, font=("Segoe UI", 11))
        pe.pack(side=tk.LEFT, padx=4)
        pe.bind("<Return>", self._phonetic_apply)
        tk.Button(pf, text="→ বাংলা", command=self._phonetic_apply,
                  bg=COLORS["accent"], fg="white", relief=tk.FLAT, padx=8, pady=2,
                  font=("Segoe UI", 9), cursor="hand2").pack(side=tk.LEFT, padx=2)

        body = tk.Frame(self, bg=COLORS["surface"])
        body.pack(padx=8, pady=6)
        for row in self.ROWS:
            rf = tk.Frame(body, bg=COLORS["surface"])
            rf.pack(pady=1)
            for ch in row:
                self._key(rf, ch, ch)

        ef = tk.Frame(body, bg=COLORS["surface"])
        ef.pack(pady=4)
        for label, val in self.EXTRA:
            self._key(ef, label, val, wide=(val in (" ", "BACK", "ENTER")))

        tk.Button(self, text="Close", command=self.destroy,
                  bg=COLORS["surface2"], fg=COLORS["text_dim"], relief=tk.FLAT,
                  padx=12, pady=4, font=("Segoe UI", 9)).pack(pady=(0, 10))

        self.geometry("+%d+%d" % (parent.winfo_rootx() + 80, parent.winfo_rooty() + 120))

    def _key(self, parent, label, value, wide=False):
        w = 6 if wide else 3
        btn = tk.Button(
            parent, text=label, width=w,
            command=lambda v=value: self._press(v),
            bg=COLORS["surface2"], fg=COLORS["text"],
            activebackground=COLORS["accent"], activeforeground="white",
            relief=tk.FLAT, font=("Segoe UI", 11), cursor="hand2", padx=2, pady=2,
        )
        btn.pack(side=tk.LEFT, padx=1, pady=1)

    def _press(self, value):
        if value == "ENTER":
            if self.on_enter:
                self.on_enter()
            return
        self.on_char(value)

    def _phonetic_apply(self, event=None):
        raw = self.phon_var.get().strip()
        if not raw:
            return
        converted = phonetic_bangla(raw)
        self.on_char(converted)
        self.phon_var.set("")


# Lightweight English-phonetic → Bangla (common patterns, longest match first)
_PHONETIC_RULES = [
    ("ksh", "ক্ষ"), ("gg", "জ্ঞ"), ("ng", "ং"), ("nj", "ঞ্জ"), ("tt", "ত্ত"),
    ("th", "থ"), ("Th", "ঠ"), ("dh", "ধ"), ("Dh", "ঢ"), ("ch", "চ"), ("Ch", "ছ"),
    ("sh", "শ"), ("Sh", "ষ"), ("ph", "ফ"), ("bh", "ভ"), ("jh", "ঝ"),
    ("kh", "খ"), ("gh", "ঘ"), ("rh", "ঢ়"),
    ("ou", "ৌ"), ("oi", "ৈ"), ("ee", "ী"), ("oo", "ূ"),
    ("aa", "া"), ("ri", "ৃ"),
    ("a", "া"), ("i", "ি"), ("I", "ী"), ("u", "ু"), ("U", "ূ"),
    ("e", "ে"), ("o", "ো"), ("O", "ৌ"),
    ("k", "ক"), ("g", "গ"), ("c", "চ"), ("j", "জ"), ("t", "ত"), ("T", "ট"),
    ("d", "দ"), ("D", "ড"), ("n", "ন"), ("N", "ণ"), ("p", "প"), ("b", "ব"),
    ("m", "ম"), ("y", "য়"), ("r", "র"), ("l", "ল"), ("s", "স"), ("h", "হ"),
    ("z", "য"), ("f", "ফ"), ("v", "ভ"), ("w", "ও"),
    ("0", "০"), ("1", "১"), ("2", "২"), ("3", "৩"), ("4", "৪"),
    ("5", "৫"), ("6", "৬"), ("7", "৭"), ("8", "৮"), ("9", "৯"),
]


def phonetic_bangla(text):
    """Convert simple phonetic English to Bangla Unicode."""
    text = text.lower()
    out = []
    i = 0
    # Independent vowels at word start-ish
    indep = {
        "a": "অ", "aa": "আ", "i": "ই", "ii": "ঈ", "u": "উ", "uu": "ঊ",
        "e": "এ", "oi": "ঐ", "o": "ও", "ou": "ঔ", "ri": "ঋ",
    }
    while i < len(text):
        if text[i] in " \t\n":
            out.append(text[i])
            i += 1
            continue
        matched = False
        # try independent vowel if at start or after space
        if i == 0 or (i > 0 and text[i - 1] in " \t\n"):
            for ln in (2, 1):
                chunk = text[i:i + ln]
                if chunk in indep:
                    out.append(indep[chunk])
                    i += ln
                    matched = True
                    break
        if matched:
            continue
        for lat, bn in _PHONETIC_RULES:
            if text.startswith(lat, i):
                out.append(bn)
                i += len(lat)
                matched = True
                break
        if not matched:
            out.append(text[i])
            i += 1
    return "".join(out)



class PDFMergerWindow(tk.Toplevel):
    """Merge multiple PDFs with order preview and Save As."""

    def __init__(self, parent):
        super().__init__(parent)
        self.parent_app = parent
        self.title("PDF Merger — One PDF Editor")
        self.configure(bg=COLORS["surface"])
        self.geometry("780x520")
        self.minsize(640, 420)
        self.transient(parent)
        self.files = []  # list of {"path", "pages", "name"}
        self._thumb = None
        self._build()

    def _build(self):
        top = tk.Frame(self, bg=COLORS["surface"])
        top.pack(fill=tk.X, padx=12, pady=10)
        tk.Label(top, text="PDF Merger", bg=COLORS["surface"], fg=COLORS["text"],
                 font=("Segoe UI", 14, "bold")).pack(side=tk.LEFT)
        tk.Label(top, text="PDF · Image · Word → auto PDF · reorder · Merge", bg=COLORS["surface"],
                 fg=COLORS["text_dim"], font=("Segoe UI", 9)).pack(side=tk.LEFT, padx=12)

        body = tk.Frame(self, bg=COLORS["surface"])
        body.pack(fill=tk.BOTH, expand=True, padx=12, pady=4)

        left = tk.Frame(body, bg=COLORS["surface"])
        left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        tk.Label(left, text="Merge order (top = first)", bg=COLORS["surface"],
                 fg=COLORS["text_dim"], font=("Segoe UI", 9)).pack(anchor=tk.W)
        lf = tk.Frame(left, bg=COLORS["surface2"])
        lf.pack(fill=tk.BOTH, expand=True, pady=4)
        self.listbox = tk.Listbox(
            lf, bg=COLORS["surface2"], fg=COLORS["text"], selectbackground=COLORS["accent"],
            font=("Segoe UI", 10), relief=tk.FLAT, activestyle="none",
            highlightthickness=0,
        )
        sb = tk.Scrollbar(lf, command=self.listbox.yview)
        self.listbox.config(yscrollcommand=sb.set)
        self.listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        self.listbox.bind("<<ListboxSelect>>", self._on_select)

        btns = tk.Frame(left, bg=COLORS["surface"])
        btns.pack(fill=tk.X, pady=6)
        for text, cmd in [
            ("+ Add files", self.add_files),
            ("↑ Up", self.move_up),
            ("↓ Down", self.move_down),
            ("Remove", self.remove_selected),
            ("Clear", self.clear_all),
        ]:
            tk.Button(
                btns, text=text, command=cmd, bg=COLORS["surface2"], fg=COLORS["text"],
                relief=tk.FLAT, padx=8, pady=4, font=("Segoe UI", 9), cursor="hand2",
            ).pack(side=tk.LEFT, padx=3)

        right = tk.Frame(body, bg=COLORS["surface"], width=260)
        right.pack(side=tk.RIGHT, fill=tk.Y, padx=(12, 0))
        right.pack_propagate(False)
        tk.Label(right, text="Preview", bg=COLORS["surface"], fg=COLORS["text_dim"],
                 font=("Segoe UI", 9)).pack(anchor=tk.W)
        self.preview_lbl = tk.Label(right, bg="#0b1220", text="No file selected",
                                    fg=COLORS["text_dim"], font=("Segoe UI", 9))
        self.preview_lbl.pack(fill=tk.BOTH, expand=True, pady=4)
        self.info_lbl = tk.Label(right, text="", bg=COLORS["surface"], fg=COLORS["text"],
                                 font=("Segoe UI", 9), justify=tk.LEFT, wraplength=240)
        self.info_lbl.pack(anchor=tk.W, pady=4)

        bottom = tk.Frame(self, bg=COLORS["surface"])
        bottom.pack(fill=tk.X, padx=12, pady=12)
        self.status = tk.Label(bottom, text="Add PDF / Image / Word files", bg=COLORS["surface"],
                               fg=COLORS["text_dim"], font=("Segoe UI", 9))
        self.status.pack(side=tk.LEFT)
        tk.Button(
            bottom, text="Merge & Save As…", command=self.merge_and_save,
            bg=COLORS["accent"], fg="white", relief=tk.FLAT, padx=16, pady=8,
            font=("Segoe UI", 10, "bold"), cursor="hand2",
        ).pack(side=tk.RIGHT)

    def add_files(self):
        paths = filedialog.askopenfilenames(
            parent=self,
            title="Add PDF / Image / Word files",
            filetypes=[
                ("All supported", "*.pdf;*.png;*.jpg;*.jpeg;*.bmp;*.gif;*.webp;*.docx"),
                ("PDF", "*.pdf"),
                ("Images", "*.png;*.jpg;*.jpeg;*.bmp;*.gif;*.webp"),
                ("Word", "*.docx"),
                ("All files", "*.*"),
            ],
        )
        for path in paths:
            try:
                kind, pages = self._probe_file(path)
                self.files.append({
                    "path": path,
                    "pages": pages,
                    "name": os.path.basename(path),
                    "kind": kind,
                })
            except Exception as e:
                messagebox.showerror("Cannot add", f"{path}\n{e}", parent=self)
        self._refresh_list()
        self.status.config(text=f"{len(self.files)} file(s) ready (PDF/Image/Word → PDF on merge)")

    def _probe_file(self, path):
        ext = Path(path).suffix.lower()
        if ext == ".pdf":
            doc = fitz.open(path)
            n = doc.page_count
            doc.close()
            return "PDF", n
        if ext in (".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp", ".tif", ".tiff"):
            return "Image", 1
        if ext == ".docx":
            if not HAS_DOCX:
                raise RuntimeError("Word support requires python-docx")
            return "Word", 1
        raise RuntimeError(f"Unsupported file type: {ext}")

    def _open_as_pdf_doc(self, path):
        """Open any supported file as a fitz PDF document (caller must close)."""
        ext = Path(path).suffix.lower()
        if ext == ".pdf":
            return fitz.open(path)
        if ext in (".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp", ".tif", ".tiff"):
            img = Image.open(path)
            if img.mode not in ("RGB", "L"):
                img = img.convert("RGB")
            bio = io.BytesIO()
            img.save(bio, format="PDF", resolution=150.0)
            bio.seek(0)
            return fitz.open("pdf", bio.read())
        if ext == ".docx":
            return docx_to_fitz(path)
        raise RuntimeError(f"Unsupported: {ext}")

    def _refresh_list(self):
        self.listbox.delete(0, tk.END)
        for i, f in enumerate(self.files, 1):
            kind = f.get("kind", "PDF")
            self.listbox.insert(tk.END, f"{i}. [{kind}]  {f['name']}  ({f['pages']} page{'s' if f['pages']!=1 else ''})")
        if self.files:
            self.listbox.selection_set(0)
            self._on_select()

    def _selected_index(self):
        sel = self.listbox.curselection()
        return int(sel[0]) if sel else None

    def move_up(self):
        i = self._selected_index()
        if i is None or i <= 0:
            return
        self.files[i - 1], self.files[i] = self.files[i], self.files[i - 1]
        self._refresh_list()
        self.listbox.selection_clear(0, tk.END)
        self.listbox.selection_set(i - 1)
        self._on_select()

    def move_down(self):
        i = self._selected_index()
        if i is None or i >= len(self.files) - 1:
            return
        self.files[i + 1], self.files[i] = self.files[i], self.files[i + 1]
        self._refresh_list()
        self.listbox.selection_clear(0, tk.END)
        self.listbox.selection_set(i + 1)
        self._on_select()

    def remove_selected(self):
        i = self._selected_index()
        if i is None:
            return
        self.files.pop(i)
        self._refresh_list()
        self.preview_lbl.config(image="", text="No file selected")
        self._thumb = None
        self.info_lbl.config(text="")
        self.status.config(text=f"{len(self.files)} file(s) in list")

    def clear_all(self):
        self.files.clear()
        self._refresh_list()
        self.preview_lbl.config(image="", text="No file selected")
        self._thumb = None
        self.info_lbl.config(text="")
        self.status.config(text="Add 2 or more PDF files")

    def _on_select(self, event=None):
        i = self._selected_index()
        if i is None or i >= len(self.files):
            return
        f = self.files[i]
        self.info_lbl.config(
            text=f"#{i + 1} in order\n[{f.get('kind','PDF')}] {f['name']}\n{f['pages']} page(s)\n(auto→PDF on merge)\n\n{f['path']}"
        )
        try:
            doc = self._open_as_pdf_doc(f["path"])
            page = doc[0]
            mat = fitz.Matrix(0.35, 0.35)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            img.thumbnail((240, 300), Image.Resampling.LANCZOS)
            self._thumb = ImageTk.PhotoImage(img)
            self.preview_lbl.config(image=self._thumb, text="")
            doc.close()
        except Exception as e:
            self.preview_lbl.config(image="", text=f"Preview failed\n{e}")
            self._thumb = None

    def merge_and_save(self):
        if len(self.files) < 1:
            messagebox.showinfo("Need files", "Add at least one PDF.", parent=self)
            return
        if len(self.files) < 2:
            if not messagebox.askyesno("One file", "Only one PDF in list. Continue anyway?", parent=self):
                return
        out = filedialog.asksaveasfilename(
            parent=self,
            title="Save merged PDF",
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf")],
            initialfile="merged.pdf",
            initialdir=str(Path.home() / "Documents"),
        )
        if not out:
            return
        if not out.lower().endswith(".pdf"):
            out += ".pdf"
        try:
            merged = fitz.open()
            total_pages = 0
            for f in self.files:
                src_doc = self._open_as_pdf_doc(f["path"])
                merged.insert_pdf(src_doc)
                total_pages += src_doc.page_count
                src_doc.close()
            merged.save(out, garbage=3, deflate=True)
            merged.close()
            self.status.config(text=f"Merged {len(self.files)} files → {total_pages} pages")
            messagebox.showinfo(
                "Done",
                f"Merged PDF saved:\n{out}\n\n{len(self.files)} files · {total_pages} pages",
                parent=self,
            )
            if messagebox.askyesno("Open?", "Open the merged PDF in One PDF Editor?", parent=self):
                try:
                    self.parent_app._load_path(out)
                except Exception:
                    pass
                self.destroy()
        except Exception as e:
            messagebox.showerror("Merge failed", str(e), parent=self)


class TabSession:
    """One open document tab (max 6)."""
    def __init__(self, title="Home"):
        self.pdf = PDFDocument()
        self.current_page = 0
        self.zoom = DEFAULT_ZOOM
        self.title = title
        self.highlight_rect = None
        self.selected_span = None
        self.selected_page = -1
        self.scroll_y = 0.0  # canvas yview fraction


class OnePDFEditor(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title(f"{APP_NAME}  v{APP_VERSION}")
        self.geometry("1200x800")
        self.minsize(900, 600)
        self.configure(bg=COLORS["bg"])
        self._set_app_icon()
        self.MAX_TABS = 6
        self.tabs = []  # list[TabSession]
        self.edit_mode = False
        self.link_edit_mode = False
        self._drop_queue = queue.Queue()
        self._drop_busy = False
        self.active_tab = -1
        self.pdf = PDFDocument()
        self.current_page = 0
        self.zoom = DEFAULT_ZOOM
        self.page_layout = []
        self.photos = []
        self.photo = None
        self.search_results = []
        self.search_index = -1
        self.highlight_rect = None
        self.selected_span = None
        self.selected_page = -1
        self._move_active = False
        self.placing_symbol = False
        self.symbol_char = None
        self._move_delta = (0.0, 0.0)
        self._move_start = None
        self.signature_img = None
        self.sig_rect = None
        self.placing_signature = False
        self._drag_start = None
        self._edit_entry = None
        self._edit_span = None
        # Color fill box tool
        self.fill_mode = False
        self.pick_color_mode = False
        self.copy_sign_mode = False
        self.screenshot_mode = False
        self.link_mode = False
        self.link_rect = None
        self.link_edit_mode = False
        self.screenshot_rect = None
        self.fill_color = None  # None = auto-sample PDF background
        self.fill_color_manual = False  # True only after Pick Color
        self.text_color = (0.0, 0.0, 0.0)  # RGB 0-1 for new text
        self.fill_rect = None
        self.copy_sign_rect = None
        self.copied_signatures = []
        self._dash_photos = []
        self._dash_idx = 0
        self._dash_job = None
        self._dash_bg_photo = None
        self._build_ui()
        self.after(100, self.render_page)
        self.after(300, self._poll_drop_queue)
        self._bind_shortcuts()
        self.protocol("WM_DELETE_WINDOW", self._on_close)
        self.report_callback_exception = self._tk_exception_guard
        self.update_idletasks()
        x = (self.winfo_screenwidth() - self.winfo_width()) // 2
        y = (self.winfo_screenheight() - self.winfo_height()) // 2
        self.geometry(f"+{x}+{y}")

    def _set_app_icon(self):
        """Set window/taskbar icon — prefer multi-size ICO on Windows, PNG fallback."""
        try:
            self._app_icon_imgs = []
            ico = resource_path("assets", "OnePDFEditor.ico")
            if ico.exists() and sys.platform.startswith("win"):
                try:
                    self.iconbitmap(default=str(ico))
                except Exception:
                    try:
                        self.iconbitmap(str(ico))
                    except Exception:
                        pass
            for name in ("icon_256.png", "icon_128.png", "icon_64.png", "icon_32.png"):
                png = resource_path("assets", name)
                if png.exists():
                    try:
                        img = Image.open(png).convert("RGBA")
                        ph = ImageTk.PhotoImage(img)
                        self._app_icon_imgs.append(ph)
                    except Exception:
                        pass
            if self._app_icon_imgs:
                try:
                    self.iconphoto(True, *self._app_icon_imgs)
                except Exception:
                    try:
                        self.iconphoto(True, self._app_icon_imgs[0])
                    except Exception:
                        pass
        except Exception:
            pass

    def _make_tool_btn(self, parent, text, command):
        return tk.Button(
            parent, text=text, command=command,
            bg=COLORS["surface2"], fg=COLORS["text"],
            activebackground=COLORS["accent_hover"], activeforeground="white",
            relief=tk.FLAT, padx=12, pady=6, cursor="hand2",
            font=("Segoe UI", 9), bd=0,
        )

    def _build_ui(self):
        menubar = tk.Menu(self, bg=COLORS["surface"], fg=COLORS["text"],
                          activebackground=COLORS["accent"], activeforeground="white", tearoff=0)
        self.config(menu=menubar)
        file_m = tk.Menu(menubar, tearoff=0, bg=COLORS["surface"], fg=COLORS["text"], activebackground=COLORS["accent"])
        menubar.add_cascade(label="File", menu=file_m)
        file_m.add_command(label="Open...\tCtrl+O", command=self.open_file)
        file_m.add_command(label="Save\tCtrl+S", command=self.save_pdf)
        file_m.add_command(label="Save As...\tCtrl+Shift+S", command=self.save_as_pdf)
        file_m.add_separator()
        file_m.add_command(label="Exit", command=self._on_close)

        edit_m = tk.Menu(menubar, tearoff=0, bg=COLORS["surface"], fg=COLORS["text"], activebackground=COLORS["accent"])
        menubar.add_cascade(label="Edit", menu=edit_m)
        edit_m.add_command(label="Undo\tCtrl+Z", command=self.undo)
        edit_m.add_command(label="Redo\tCtrl+Y", command=self.redo)
        edit_m.add_command(label="Search...\tCtrl+F", command=self.show_search)

        view_m = tk.Menu(menubar, tearoff=0, bg=COLORS["surface"], fg=COLORS["text"], activebackground=COLORS["accent"])
        menubar.add_cascade(label="View", menu=view_m)
        view_m.add_command(label="Zoom In", command=lambda: self.set_zoom(self.zoom * 1.25))
        view_m.add_command(label="Zoom Out", command=lambda: self.set_zoom(self.zoom / 1.25))
        view_m.add_command(label="Fit Page", command=self.fit_page)
        view_m.add_command(label="Fit Width", command=self.fit_width)
        view_m.add_command(label="Actual Size (100%)", command=lambda: self.set_zoom(1.0))

        tools_m = tk.Menu(menubar, tearoff=0, bg=COLORS["surface"], fg=COLORS["text"], activebackground=COLORS["accent"])
        menubar.add_cascade(label="Tools", menu=tools_m)
        tools_m.add_command(label="Draw Signature...", command=self.start_signature)
        tools_m.add_command(label="Copy Sign (select area)...", command=self.start_copy_sign_region)
        tools_m.add_separator()
        tools_m.add_command(label="Color Fill Box (Hide Text)...", command=self.start_fill_box)
        tools_m.add_command(label="Pick Fill Color from Page...", command=self.start_pick_color)
        tools_m.add_command(label="Reset Fill Color (auto background)", command=self.reset_fill_color)
        tools_m.add_separator()
        tools_m.add_command(label="Text Color...", command=self.choose_text_color)
        tools_m.add_command(label="বাংলা কিবোর্ড...", command=self._open_bangla_kb_standalone)
        tools_m.add_separator()
        tools_m.add_command(label="Screenshot", command=self.take_screenshot)
        tools_m.add_command(label="OCR Image Text...", command=self.ocr_current_page)
        tools_m.add_separator()
        tools_m.add_command(label="PDF Merger...", command=self.open_pdf_merger)
        tools_m.add_separator()
        tools_m.add_command(label="Hyperlinks on Page...", command=self.manage_hyperlinks)
        tools_m.add_command(label="Add Hyperlink...", command=self.start_add_hyperlink)
        tools_m.add_command(label="Add Link to Selected Text...", command=self.add_link_to_selected_text)
        tools_m.add_command(label="Set as Default PDF Viewer...", command=self.show_default_viewer_help)

        help_m = tk.Menu(menubar, tearoff=0, bg=COLORS["surface"], fg=COLORS["text"], activebackground=COLORS["accent"])
        sym_m = tk.Menu(menubar, tearoff=0, bg=COLORS["surface"], fg=COLORS["text"], activebackground=COLORS["accent"])
        menubar.add_cascade(label="Symbols", menu=sym_m)
        for lab, ch in [
            ("✓  Check mark", "✓"),
            ("✔  Heavy check", "✔"),
            ("☑  Ballot box check", "☑"),
            ("☐  Empty box", "☐"),
            ("✗  Cross", "✗"),
            ("★  Star", "★"),
            ("•  Bullet", "•"),
            ("→  Arrow", "→"),
            ("©  Copyright", "©"),
            ("®  Registered", "®"),
            ("°  Degree", "°"),
            ("§  Section", "§"),
            ("—  Em dash", "—"),
            ("€  Euro", "€"),
            ("£  Pound", "£"),
            ("₹  Rupee", "₹"),
            ("™  Trademark", "™"),
            ("∞  Infinity", "∞"),
            ("≈  Approx", "≈"),
            ("±  Plus-minus", "±"),
        ]:
            sym_m.add_command(label=lab, command=lambda c=ch: self.start_place_symbol(c))
        menubar.add_cascade(label="Help", menu=help_m)
        help_m.add_command(label="About", command=self.show_about)

        toolbar = tk.Frame(self, bg=COLORS["toolbar"])
        toolbar.pack(side=tk.TOP, fill=tk.X)
        # Row 1 — file / edit / tools (always visible)
        row1 = tk.Frame(toolbar, bg=COLORS["toolbar"])
        row1.pack(side=tk.TOP, fill=tk.X, padx=6, pady=(6, 2))
        for text, cmd in [("Open", self.open_file), ("Save", self.save_pdf), ("Save As", self.save_as_pdf)]:
            self._make_tool_btn(row1, text, cmd).pack(side=tk.LEFT, padx=2)
        self.btn_edit_mode = self._make_tool_btn(row1, "View Mode", self.toggle_edit_mode)
        self.btn_edit_mode.pack(side=tk.LEFT, padx=6)
        self.btn_edit_mode.config(bg="#0ea5e9", fg="white")
        self._make_tool_btn(row1, "Print", self.print_document).pack(side=tk.LEFT, padx=2)
        self._make_tool_btn(row1, "Merge", self.open_pdf_merger).pack(side=tk.LEFT, padx=2)
        tk.Frame(row1, width=8, bg=COLORS["toolbar"]).pack(side=tk.LEFT)
        for text, cmd in [("Undo", self.undo), ("Search", self.show_search)]:
            self._make_tool_btn(row1, text, cmd).pack(side=tk.LEFT, padx=2)
        self._make_tool_btn(row1, "✓", lambda: self.start_place_symbol("✓")).pack(side=tk.LEFT, padx=2)
        tk.Frame(row1, width=8, bg=COLORS["toolbar"]).pack(side=tk.LEFT)
        for text, cmd in [("Sign", self.start_signature), ("Copy Sign", self.start_copy_sign_region),
                          ("Fill Box", self.start_fill_box), ("Screenshot", self.take_screenshot)]:
            self._make_tool_btn(row1, text, cmd).pack(side=tk.LEFT, padx=2)
        # Row 2 — page / zoom (fits small widths)
        row2 = tk.Frame(toolbar, bg=COLORS["toolbar"])
        row2.pack(side=tk.TOP, fill=tk.X, padx=6, pady=(2, 6))
        self._make_tool_btn(row2, "Prev", self.prev_page).pack(side=tk.LEFT, padx=2)
        self._make_tool_btn(row2, "Next", self.next_page).pack(side=tk.LEFT, padx=2)
        tk.Label(row2, text="Page", bg=COLORS["toolbar"], fg=COLORS["text_dim"], font=("Segoe UI", 9)).pack(side=tk.LEFT, padx=(8, 2))
        self.page_var = tk.StringVar(value="1")
        page_entry = tk.Entry(row2, textvariable=self.page_var, width=4, bg=COLORS["surface2"], fg=COLORS["text"],
                              insertbackground=COLORS["text"], relief=tk.FLAT, font=("Segoe UI", 10))
        page_entry.pack(side=tk.LEFT)
        page_entry.bind("<Return>", self._goto_page)
        self.page_count_lbl = tk.Label(row2, text="/ 0", bg=COLORS["toolbar"], fg=COLORS["text_dim"], font=("Segoe UI", 9))
        self.page_count_lbl.pack(side=tk.LEFT, padx=4)
        tk.Frame(row2, width=12, bg=COLORS["toolbar"]).pack(side=tk.LEFT)
        self._make_tool_btn(row2, "-", lambda: self.set_zoom(self.zoom / 1.25)).pack(side=tk.LEFT)
        self.zoom_lbl = tk.Label(row2, text="100%", bg=COLORS["toolbar"], fg=COLORS["text"], font=("Segoe UI", 9), width=5)
        self.zoom_lbl.pack(side=tk.LEFT, padx=4)
        self._make_tool_btn(row2, "+", lambda: self.set_zoom(self.zoom * 1.25)).pack(side=tk.LEFT)
        self._make_tool_btn(row2, "Fit", self.fit_page).pack(side=tk.LEFT, padx=4)

        # Tab bar (up to 6 open files)
        self.tab_bar = tk.Frame(self, bg=COLORS["surface"], height=34)
        self.tab_bar.pack(side=tk.TOP, fill=tk.X)
        self.tab_bar.pack_propagate(False)
        self._tab_buttons = []
        self._refresh_tab_bar()

        self.drop_hint = tk.Label(self, text="  Open PDF / Image / Word file  •  Double-click text to edit in place",
                                  bg=COLORS["surface2"], fg=COLORS["text_dim"], font=("Segoe UI", 10), pady=8)
        self.drop_hint.pack(side=tk.TOP, fill=tk.X)

        main = tk.Frame(self, bg=COLORS["canvas_bg"])
        main.pack(fill=tk.BOTH, expand=True)
        self.canvas = tk.Canvas(main, bg=COLORS["canvas_bg"], highlightthickness=0)
        hsb = ttk.Scrollbar(main, orient=tk.HORIZONTAL, command=self.canvas.xview)
        vsb = ttk.Scrollbar(main, orient=tk.VERTICAL, command=self.canvas.yview)
        self.canvas.configure(xscrollcommand=hsb.set, yscrollcommand=vsb.set)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        hsb.pack(side=tk.BOTTOM, fill=tk.X)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.canvas.bind("<Button-1>", self._on_canvas_click)
        self.canvas.bind("<Double-Button-1>", self._on_double_click)
        self.canvas.bind("<B1-Motion>", self._on_canvas_drag)
        self.canvas.bind("<ButtonRelease-1>", self._on_canvas_release)
        self.canvas.bind("<Configure>", self._on_canvas_configure)
        self.canvas.bind("<MouseWheel>", self._on_mousewheel)
        # Drag-drop: windnd is Windows-only
        if sys.platform.startswith("win"):
            try:
                import windnd
                windnd.hook_dropfiles(self, func=self._on_windnd_drop)
            except Exception:
                pass
        self.canvas.bind("<Button-4>", lambda e: self.canvas.yview_scroll(-1, "units"))
        self.canvas.bind("<Button-5>", lambda e: self.canvas.yview_scroll(1, "units"))

        self.status = tk.Label(self, text=f"{APP_NAME} v{APP_VERSION}  •  Made by {APP_STUDIO}",
                               bg=COLORS["surface"], fg=COLORS["text_dim"], font=("Segoe UI", 9),
                               anchor=tk.W, padx=10, pady=4)
        self.status.pack(side=tk.BOTTOM, fill=tk.X)

        self.search_frame = tk.Frame(self, bg=COLORS["surface2"])
        tk.Label(self.search_frame, text="Find:", bg=COLORS["surface2"], fg=COLORS["text"], font=("Segoe UI", 10)).pack(side=tk.LEFT, padx=8)
        self.search_var = tk.StringVar()
        self.search_entry = tk.Entry(self.search_frame, textvariable=self.search_var, width=28,
                                     bg=COLORS["surface"], fg=COLORS["text"], insertbackground=COLORS["text"],
                                     relief=tk.FLAT, font=("Segoe UI", 10))
        self.search_entry.pack(side=tk.LEFT, padx=4, pady=6)
        self.search_entry.bind("<Return>", lambda e: self.do_search())
        self._make_tool_btn(self.search_frame, "Next", self.find_next).pack(side=tk.LEFT, padx=2)
        self._make_tool_btn(self.search_frame, "Prev", self.find_prev).pack(side=tk.LEFT, padx=2)
        self._make_tool_btn(self.search_frame, "Close", self.hide_search).pack(side=tk.LEFT, padx=6)
        self.search_status = tk.Label(self.search_frame, text="", bg=COLORS["surface2"], fg=COLORS["text_dim"], font=("Segoe UI", 9))
        self.search_status.pack(side=tk.LEFT, padx=8)

    def _bind_shortcuts(self):
        self.bind("<Control-o>", lambda e: self.open_file())
        self.bind("<Control-p>", lambda e: self.print_document())
        self.bind("<Control-w>", lambda e: self.close_tab())
        self.bind("<Control-W>", lambda e: self.close_tab())
        self.bind("<Control-P>", lambda e: self.print_document())
        self.bind("<Control-s>", lambda e: self.save_pdf())
        self.bind("<Control-S>", lambda e: self.save_as_pdf())
        self.bind("<Control-Shift-S>", lambda e: self.save_as_pdf())
        self.bind("<Control-Shift-s>", lambda e: self.save_as_pdf())
        self.bind("<Control-f>", lambda e: self.show_search())
        self.bind("<Control-z>", lambda e: self.undo())
        self.bind("<Control-y>", lambda e: self.redo())
        # macOS Command key
        self.bind("<Command-o>", lambda e: self.open_file())
        self.bind("<Command-p>", lambda e: self.print_document())
        self.bind("<Command-w>", lambda e: self.close_tab())
        self.bind("<Command-s>", lambda e: self.save_pdf())
        self.bind("<Command-Shift-s>", lambda e: self.save_as_pdf())
        self.bind("<Command-f>", lambda e: self.show_search())
        self.bind("<Command-z>", lambda e: self.undo())
        self.bind("<Command-y>", lambda e: self.redo())
        self.bind("<Escape>", lambda e: self._cancel_ops())
        self.bind("<Prior>", lambda e: self.prev_page())
        self.bind("<Next>", lambda e: self.next_page())
        self.bind("<Left>", lambda e: self.prev_page())
        self.bind("<Right>", lambda e: self.next_page())
        self.bind("<Up>", lambda e: self.prev_page())
        self.bind("<Down>", lambda e: self.next_page())


    def _push_ui_to_tab(self):
        if self.active_tab < 0 or self.active_tab >= len(self.tabs):
            return
        # Capture scroll-synced page before saving
        try:
            self._sync_page_from_scroll()
        except Exception:
            pass
        t = self.tabs[self.active_tab]
        t.pdf = self.pdf
        t.current_page = int(self.current_page)
        t.zoom = float(self.zoom)
        t.highlight_rect = self.highlight_rect
        t.selected_span = self.selected_span
        t.selected_page = self.selected_page
        try:
            t.scroll_y = float(self.canvas.yview()[0])
        except Exception:
            t.scroll_y = 0.0
        if self.pdf.path:
            t.title = os.path.basename(self.pdf.path)
        elif self.pdf.doc:
            t.title = t.title or "Untitled"

    def _pull_tab_to_ui(self):
        if self.active_tab < 0 or self.active_tab >= len(self.tabs):
            self.pdf = PDFDocument()
            self.current_page = 0
            self.zoom = DEFAULT_ZOOM
            self.highlight_rect = None
            self.selected_span = None
            self.selected_page = -1
            return
        t = self.tabs[self.active_tab]
        self.pdf = t.pdf
        # Clamp page to this document's range
        n = t.pdf.page_count() if t.pdf.doc else 0
        cp = int(t.current_page or 0)
        if n > 0:
            cp = max(0, min(cp, n - 1))
        else:
            cp = 0
        self.current_page = cp
        t.current_page = cp
        self.zoom = float(t.zoom or DEFAULT_ZOOM)
        self.highlight_rect = t.highlight_rect
        self.selected_span = t.selected_span
        self.selected_page = t.selected_page
        try:
            self.zoom_lbl.config(text=f"{int(self.zoom * 100)}%")
        except Exception:
            pass

    def _refresh_tab_bar(self):
        for w in getattr(self, "_tab_buttons", []):
            try:
                w.destroy()
            except Exception:
                pass
        self._tab_buttons = []
        if not hasattr(self, "tab_bar"):
            return
        for i, t in enumerate(self.tabs):
            active = i == self.active_tab
            bg = COLORS["accent"] if active else COLORS["surface2"]
            fg = "white" if active else COLORS["text"]
            name = t.title or "Untitled"
            if len(name) > 18:
                name = name[:15] + "…"
            dirty = " •" if getattr(t.pdf, "dirty", False) else ""
            fr = tk.Frame(self.tab_bar, bg=bg)
            fr.pack(side=tk.LEFT, padx=(4, 0), pady=4)
            btn = tk.Button(
                fr, text=f"{name}{dirty}", command=lambda idx=i: self.switch_tab(idx),
                bg=bg, fg=fg, relief=tk.FLAT, font=("Segoe UI", 9),
                padx=10, pady=2, cursor="hand2", activebackground=COLORS["accent_hover"],
            )
            btn.pack(side=tk.LEFT)
            close = tk.Button(
                fr, text="×", command=lambda idx=i: self.close_tab(idx),
                bg=bg, fg=fg, relief=tk.FLAT, font=("Segoe UI", 10, "bold"),
                padx=6, pady=2, cursor="hand2", activebackground="#ef4444",
            )
            close.pack(side=tk.LEFT)
            self._tab_buttons.extend([fr, btn, close])
        if not self.tabs:
            empty = tk.Label(
                self.tab_bar, text="No files open — Open or drag a file (max 6 tabs)",
                bg=COLORS["surface"], fg=COLORS["text_dim"], font=("Segoe UI", 9),
            )
            empty.pack(side=tk.LEFT, padx=10, pady=6)
            self._tab_buttons.append(empty)

    def switch_tab(self, idx):
        if idx < 0 or idx >= len(self.tabs):
            return
        if idx == self.active_tab:
            return
        self._cancel_ops()
        self._push_ui_to_tab()
        self.active_tab = idx
        self._pull_tab_to_ui()
        self._refresh_tab_bar()
        self._update_page_ui()
        self.render_page()
        # Restore this tab's page / scroll (must be after render builds layout)
        def _restore():
            try:
                t = self.tabs[self.active_tab]
                n = self.pdf.page_count() if self.pdf.doc else 0
                if n > 0:
                    self.current_page = max(0, min(int(t.current_page), n - 1))
                    self._update_page_ui()
                    # Prefer exact scroll fraction if valid, else jump to saved page
                    sy = float(getattr(t, "scroll_y", 0.0) or 0.0)
                    if 0.0 <= sy <= 1.0 and sy > 0.001:
                        self.canvas.yview_moveto(sy)
                    else:
                        self.scroll_to_page(self.current_page)
            except Exception:
                pass
        self.after(20, _restore)
        if self.pdf.doc:
            try:
                self.drop_hint.pack_forget()
            except Exception:
                pass
            self.status.config(
                text=f"Tab {idx + 1}: {self.tabs[idx].title}  ·  page {self.current_page + 1}"
            )
        else:
            self.status.config(text="Empty tab")

    def close_tab(self, idx=None):
        if idx is None:
            idx = self.active_tab
        if idx < 0 or idx >= len(self.tabs):
            return
        t = self.tabs[idx]
        if t.pdf.dirty:
            if not messagebox.askyesno("Unsaved", f'"{t.title}" has unsaved changes. Close anyway?', parent=self):
                return
        try:
            t.pdf.close()
        except Exception:
            pass
        was_active = (idx == self.active_tab)
        self.tabs.pop(idx)
        if not self.tabs:
            self.active_tab = -1
            self.pdf = PDFDocument()
            self.current_page = 0
            self.zoom = DEFAULT_ZOOM
            self.highlight_rect = None
            self.selected_span = None
            self.selected_page = -1
            self.page_layout = []
            self.photos = []
            self._refresh_tab_bar()
            self._update_page_ui()
            self.render_page()
            try:
                self.drop_hint.pack(side=tk.TOP, fill=tk.X)
            except Exception:
                pass
            self.status.config(text="All tabs closed")
            return
        # Adjust active index; never close sibling tabs' documents
        if was_active:
            self.active_tab = min(idx, len(self.tabs) - 1)
            self._pull_tab_to_ui()
        else:
            if idx < self.active_tab:
                self.active_tab -= 1
            # active document object unchanged
        self._refresh_tab_bar()
        self._update_page_ui()
        self.render_page()
        self.status.config(text=f"Closed tab · {len(self.tabs)} still open")

    def _open_in_tab(self, path):
        """Create a NEW tab for path (caller ensures not already open). Max 6."""
        if len(self.tabs) >= self.MAX_TABS:
            messagebox.showwarning(
                "Tab limit",
                f"Maximum {self.MAX_TABS} tabs open. Close a tab first.",
                parent=self,
            )
            return False
        self._push_ui_to_tab()
        title = os.path.basename(path) if path else "Untitled"
        session = TabSession(title=title)
        self.tabs.append(session)
        self.active_tab = len(self.tabs) - 1
        self.pdf = session.pdf
        self.current_page = 0
        self.zoom = DEFAULT_ZOOM
        self.highlight_rect = None
        self.selected_span = None
        self.selected_page = -1
        self.search_results = []
        self.search_index = -1
        return True


    def open_file(self):
        # Multi-tab: opening a new file does NOT close other tabs
        if len(self.tabs) >= self.MAX_TABS:
            messagebox.showwarning(
                "Tab limit",
                f"Already {self.MAX_TABS} tabs open.\nClose a tab (×) before opening another.",
                parent=self,
            )
            return
        path = filedialog.askopenfilename(
            title="Open File",
            filetypes=[
                ("All supported", "*.pdf;*.png;*.jpg;*.jpeg;*.bmp;*.gif;*.webp;*.docx"),
                ("PDF", "*.pdf"),
                ("Images", "*.png;*.jpg;*.jpeg;*.bmp;*.gif;*.webp;*.tiff"),
                ("Word", "*.docx"),
                ("All files", "*.*"),
            ],
        )
        if path:
            self._load_path(path)

    def _load_path(self, path):
        try:
            self._stop_dashboard()
            self._cancel_ops()
            path = str(path)
            # Already open in a tab? just switch (do not reload / do not close others)
            try:
                resolved = str(Path(path).resolve())
            except Exception:
                resolved = path
            for i, t in enumerate(self.tabs):
                try:
                    if t.pdf.path and str(Path(t.pdf.path).resolve()) == resolved and t.pdf.doc:
                        self.switch_tab(i)
                        self.status.config(text=f"Switched to open tab: {t.title}")
                        return
                except Exception:
                    pass
            if not self._open_in_tab(path):
                return
            ok = self.pdf.open(path)
            if not ok:
                pwd = simpledialog.askstring("Password", "File is encrypted. Enter password:", show="*", parent=self)
                if pwd is None:
                    # rollback empty tab
                    self.close_tab(self.active_tab)
                    return
                ok = self.pdf.open(path, password=pwd)
                if not ok:
                    messagebox.showerror("Error", "Incorrect password or cannot open.", parent=self)
                    self.close_tab(self.active_tab)
                    return
            self.current_page = 0
            self.zoom = DEFAULT_ZOOM
            self.highlight_rect = None
            self.selected_span = None
            self._push_ui_to_tab()
            self._refresh_tab_bar()
            self.set_edit_mode(False)
            self._update_page_ui()
            self.render_page()
            self.fit_width()
            name = os.path.basename(path)
            kind = self.pdf.source_type.upper()
            n = self.pdf.page_count()
            self.status.config(text=f"Opened ({kind}): {name}  ·  {n} page(s)  ·  Tab {self.active_tab + 1}/{len(self.tabs)}")
            try:
                self.drop_hint.pack_forget()
            except Exception:
                pass
        except Exception as e:
            messagebox.showerror("Open failed", str(e), parent=self)
            try:
                if self.active_tab >= 0 and not self.pdf.doc:
                    self.close_tab(self.active_tab)
            except Exception:
                pass

    def save_pdf(self):
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a file first.", parent=self)
            return
        if not self.pdf.path or not str(self.pdf.path).lower().endswith(".pdf"):
            return self.save_as_pdf()
        try:
            self.pdf.save()
            self.status.config(text=f"Saved: {os.path.basename(self.pdf.path)}")
            messagebox.showinfo("Saved", f"Saved to:\n{self.pdf.path}", parent=self)
        except Exception as e:
            messagebox.showerror("Save failed", str(e), parent=self)

    def save_as_pdf(self):
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a file first.", parent=self)
            return
        # Ensure window is focused so native dialog appears on Windows
        try:
            self.lift()
            self.focus_force()
            self.update_idletasks()
        except Exception:
            pass
        initial_name = Path(self.pdf.path or "document").stem + ".pdf"
        initial_dir = None
        try:
            if self.pdf.path and Path(self.pdf.path).parent.exists():
                initial_dir = str(Path(self.pdf.path).parent)
            else:
                initial_dir = str(Path.home() / "Documents")
        except Exception:
            initial_dir = str(Path.home())
        path = filedialog.asksaveasfilename(
            parent=self,
            title="Save As — choose name and folder",
            defaultextension=".pdf",
            filetypes=[
                ("PDF files", "*.pdf"),
                ("All files", "*.*"),
            ],
            initialfile=initial_name,
            initialdir=initial_dir,
            confirmoverwrite=True,
        )
        if not path:
            self.status.config(text="Save As cancelled")
            return
        if not path.lower().endswith(".pdf"):
            path = path + ".pdf"
        try:
            self.pdf.save(path)
            self.status.config(text=f"Saved as: {os.path.basename(path)}")
            messagebox.showinfo("Saved", f"File saved as:\n{path}", parent=self)
        except Exception as e:
            messagebox.showerror("Save As failed", str(e), parent=self)

    def undo(self):
        self._cancel_inline_edit()
        if self.pdf.undo():
            self.highlight_rect = None
            self.selected_span = None
            self.render_page()
            self.status.config(text="Undo")
        else:
            self.status.config(text="Nothing to undo")

    def redo(self):
        self._cancel_inline_edit()
        if self.pdf.redo():
            self.highlight_rect = None
            self.selected_span = None
            self.render_page()
            self.status.config(text="Redo")
        else:
            self.status.config(text="Nothing to redo")

    def prev_page(self, event=None):
        if not self.pdf.doc:
            return
        if self.current_page > 0:
            self._cancel_inline_edit()
            self._move_active = False
            self.current_page -= 1
            self.highlight_rect = None
            self.selected_span = None
            self._update_page_ui()
            self.render_page()
            self.scroll_to_page(self.current_page)
        else:
            self.status.config(text="Already on first page")

    def next_page(self, event=None):
        if not self.pdf.doc:
            return
        n = self.pdf.page_count()
        if self.current_page < n - 1:
            self._cancel_inline_edit()
            self._move_active = False
            self.current_page += 1
            self.highlight_rect = None
            self.selected_span = None
            self._update_page_ui()
            self.render_page()
            self.scroll_to_page(self.current_page)
        else:
            self.status.config(text=f"Already on last page ({n})")

    def _goto_page(self, event=None):
        try:
            p = int(self.page_var.get()) - 1
            if 0 <= p < self.pdf.page_count():
                self._cancel_inline_edit()
                self.current_page = p
                self.highlight_rect = None
                self.selected_span = None
                self._update_page_ui()
                self.render_page()
        except ValueError:
            pass

    def set_zoom(self, z):
        self._cancel_inline_edit()
        self.zoom = max(0.25, min(4.0, z))
        self.zoom_lbl.config(text=f"{int(self.zoom * 100)}%")
        self.render_page()

    def fit_page(self):
        if not self.pdf.doc:
            return
        page = self.pdf.get_page(self.current_page)
        if not page:
            return
        cw = max(self.canvas.winfo_width(), 100)
        ch = max(self.canvas.winfo_height(), 100)
        rect = page.rect
        self.set_zoom(min((cw - 24) / rect.width, (ch - 24) / rect.height))

    def fit_width(self):
        if not self.pdf.doc:
            return
        page = self.pdf.get_page(self.current_page)
        if not page:
            return
        cw = max(self.canvas.winfo_width(), 100)
        self.set_zoom((cw - 24) / page.rect.width)

    def _update_page_ui(self):
        n = self.pdf.page_count() if self.pdf.doc else 0
        cur = (self.current_page + 1) if n else 0
        if self.current_page >= n and n > 0:
            self.current_page = n - 1
            cur = n
        self.page_count_lbl.config(text=f"/ {n}")
        self.page_var.set(str(cur) if n else "-")
        try:
            name = os.path.basename(self.pdf.path) if self.pdf.path else ""
            if n:
                self.status.config(text=f"Page {cur} of {n}" + (f"  ·  {name}" if name else ""))
        except Exception:
            pass

    def _on_mousewheel(self, event):
        if not self.pdf.doc and not getattr(self, "_dash_photos", None):
            return
        delta = 0
        if getattr(event, "delta", 0):
            delta = int(-1 * (event.delta / 120))
        elif getattr(event, "num", None) == 4:
            delta = -1
        elif getattr(event, "num", None) == 5:
            delta = 1
        if not delta:
            return
        # Ctrl + wheel = zoom
        try:
            if event.state & 0x0004:
                self.set_zoom(self.zoom * (1.1 if delta < 0 else (1 / 1.1)))
                return
        except Exception:
            pass
        self.canvas.yview_scroll(delta, "units")
        # Update current page from scroll position
        try:
            self._sync_page_from_scroll()
        except Exception:
            pass


    def _sync_page_from_scroll(self):
        if not self.page_layout:
            return
        # top of visible area
        try:
            top_frac = self.canvas.yview()[0]
            y = top_frac * float(self.canvas.bbox("all")[3])
        except Exception:
            return
        for pl in self.page_layout:
            if pl["y0"] <= y < pl["y1"]:
                if self.current_page != pl["idx"]:
                    self.current_page = pl["idx"]
                    self._update_page_ui()
                return

    def scroll_to_page(self, page_idx):
        for pl in self.page_layout:
            if pl["idx"] == page_idx:
                try:
                    total = float(self.canvas.bbox("all")[3]) or 1.0
                    self.canvas.yview_moveto(max(0.0, pl["y0"] / total))
                except Exception:
                    pass
                return

    def render_page(self):
        """Render ALL pages stacked vertically for continuous mouse scroll."""
        self.canvas.delete("all")
        self.photo = None
        self.photos = []
        self.page_layout = []
        if not self.pdf.doc:
            self._show_dashboard()
            return
        self._stop_dashboard()
        try:
            mat = fitz.Matrix(self.zoom, self.zoom)
            y = 0.0
            gap = 16
            max_w = 0
            n = self.pdf.page_count()
            for i in range(n):
                page = self.pdf.get_page(i)
                if not page:
                    continue
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                photo = ImageTk.PhotoImage(img)
                self.photos.append(photo)
                # centered-ish left margin 0
                self.canvas.create_image(0, y, anchor=tk.NW, image=photo, tags=("page", f"p{i}"))
                # subtle page separator label
                self.canvas.create_text(
                    8, y + 6, anchor=tk.NW, text=f"Page {i + 1} / {n}",
                    fill="#64748b", font=("Segoe UI", 8), tags=("pagelabel", f"p{i}"),
                )
                h = float(pix.height)
                w = float(pix.width)
                self.page_layout.append({"idx": i, "y0": y, "y1": y + h, "w": w, "h": h})
                max_w = max(max_w, w)
                y += h + gap
            total_h = max(y - gap, 1)
            self.canvas.config(scrollregion=(0, 0, max_w, total_h), bg="#1a1a24")

            # Overlays for current interaction page
            def canvas_rect(page_idx, rect):
                off = 0.0
                for pl in self.page_layout:
                    if pl["idx"] == page_idx:
                        off = pl["y0"]
                        break
                r = fitz.Rect(rect) * self.zoom
                return fitz.Rect(r.x0, r.y0 + off, r.x1, r.y1 + off)

            if self.highlight_rect is not None:
                r = canvas_rect(self.current_page, self.highlight_rect)
                self.canvas.create_rectangle(r.x0, r.y0, r.x1, r.y1, outline="#fbbf24", width=2, tags="highlight")
            if self.selected_span and self.selected_page is not None:
                r0 = fitz.Rect(self.selected_span["bbox"])
                if self._move_active and self._move_delta:
                    r0 = fitz.Rect(
                        r0.x0 + self._move_delta[0], r0.y0 + self._move_delta[1],
                        r0.x1 + self._move_delta[0], r0.y1 + self._move_delta[1],
                    )
                r = canvas_rect(self.selected_page, r0)
                self.canvas.create_rectangle(r.x0, r.y0, r.x1, r.y1, outline="#4ade80", width=2, tags="selected")
            if self.placing_signature and self.sig_rect:
                r = canvas_rect(self.current_page, self.sig_rect)
                self.canvas.create_rectangle(r.x0, r.y0, r.x1, r.y1, outline="#7c5cff", width=2, dash=(4, 2), tags="sigpreview")
            if self.fill_mode and self.fill_rect:
                r = canvas_rect(self.current_page, self.fill_rect)
                if self.fill_color_manual and self.fill_color:
                    c = self.fill_color
                    fill_hex = "#%02x%02x%02x" % (int(c[0] * 255), int(c[1] * 255), int(c[2] * 255))
                else:
                    fill_hex = "#cccccc"
                self.canvas.create_rectangle(r.x0, r.y0, r.x1, r.y1, outline="#f87171", width=2,
                                            fill=fill_hex, stipple="gray50", tags="fillpreview")
            if self.copy_sign_mode and self.copy_sign_rect:
                r = canvas_rect(self.current_page, self.copy_sign_rect)
                self.canvas.create_rectangle(r.x0, r.y0, r.x1, r.y1, outline="#22d3ee", width=2,
                                            dash=(4, 2), tags="copysignpreview")
            if self.screenshot_mode and self.screenshot_rect:
                r = canvas_rect(self.current_page, self.screenshot_rect)
                self.canvas.create_rectangle(r.x0, r.y0, r.x1, r.y1, outline="#fbbf24", width=2,
                                            dash=(5, 3), tags="sspreview")
            if self.link_mode and self.link_rect:
                r = canvas_rect(self.current_page, self.link_rect)
                self.canvas.create_rectangle(r.x0, r.y0, r.x1, r.y1, outline="#38bdf8", width=2,
                                            dash=(4, 2), tags="linkpreview")
            if self.link_edit_mode:
                try:
                    for lk in self.pdf.get_page_links(self.current_page):
                        r = canvas_rect(self.current_page, lk.get("from"))
                        self.canvas.create_rectangle(
                            r.x0, r.y0, r.x1, r.y1, outline="#38bdf8", width=2,
                            dash=(3, 2), tags="link_highlight",
                        )
                except Exception:
                    pass
        except Exception as e:
            self.status.config(text=f"Render error: {e}")

    def _canvas_to_pdf(self, x, y):
        """Map canvas coords to PDF page coords; updates current_page from Y position."""
        page_idx = self.current_page
        local_y = y
        for pl in self.page_layout:
            if pl["y0"] <= y <= pl["y1"] + 8:
                page_idx = pl["idx"]
                local_y = y - pl["y0"]
                break
        else:
            # nearest page
            if self.page_layout:
                pl = min(self.page_layout, key=lambda p: abs((p["y0"] + p["y1"]) / 2 - y))
                page_idx = pl["idx"]
                local_y = y - pl["y0"]
        self.current_page = page_idx
        try:
            self.page_var.set(str(page_idx + 1))
        except Exception:
            pass
        return fitz.Point(x / self.zoom, local_y / self.zoom)

    def _on_canvas_click(self, event):
        if not self.pdf.doc:
            return
        if self._edit_entry:
            self._finish_inline_edit()
            return
        cx = self.canvas.canvasx(event.x)
        cy = self.canvas.canvasy(event.y)
        pt = self._canvas_to_pdf(cx, cy)

        # Place symbol
        if self.placing_symbol and self.symbol_char:
            self._place_symbol_at(pt)
            return

        # Color picker mode
        if self.pick_color_mode:
            self._sample_color_at(cx, cy)
            return

        # Screenshot select area
        if self.screenshot_mode:
            self._drag_start = pt
            self.screenshot_rect = fitz.Rect(pt, pt)
            return

        # Copy sign region mode
        if self.copy_sign_mode:
            self._drag_start = pt
            self.copy_sign_rect = fitz.Rect(pt, pt)
            return

        # Add hyperlink mode – start drag
        if self.link_mode:
            self._drag_start = pt
            self.link_rect = fitz.Rect(pt, pt)
            return

        # Fill box mode – start drag
        if self.fill_mode:
            self._drag_start = pt
            self.fill_rect = fitz.Rect(pt, pt)
            return

        if self.placing_signature and self.signature_img:
            self._drag_start = pt
            self.sig_rect = fitz.Rect(pt, pt)
            return
        tool_busy = any([
            self.link_mode, self.fill_mode, self.screenshot_mode, self.copy_sign_mode,
            self.placing_signature, self.placing_symbol, getattr(self, "pick_color_mode", False),
        ])

        # Link-edit mode: click a link area → edit URL dialog
        if self.link_edit_mode and not tool_busy:
            if self._edit_link_at(self.current_page, pt):
                return

        # VIEW mode: click link → browser (text stays unchanged)
        if (not self.edit_mode) and (not tool_busy):
            if self._open_link_at(self.current_page, pt):
                return
            # no text selection in view mode
            self.selected_span = None
            self.render_page()
            self.status.config(text="VIEW mode — click a link to open · switch to Edit Mode to change content")
            return

        # EDIT mode: normal tools / text select
        spans = self.pdf.get_text_spans(self.current_page)
        hit = None
        for s in spans:
            if s["bbox"].contains(pt):
                hit = s
                break
        if hit:
            self.selected_span = hit
            self.selected_page = self.current_page
            self._move_active = True
            self._move_start = pt
            self._move_delta = (0.0, 0.0)
            self._drag_start = pt
            self.status.config(text=f'Selected: "{hit["text"][:40]}"  —  Drag to move · Double-click to edit')
            try:
                self.canvas.config(cursor="fleur")
            except Exception:
                pass
            self.render_page()
        else:
            self.selected_span = None
            self._move_active = False
            self._move_start = None
            self._move_delta = (0.0, 0.0)
            try:
                self.canvas.config(cursor="")
            except Exception:
                pass
            self.render_page()

    def _on_double_click(self, event):
        if not self.edit_mode:
            self.status.config(text="Switch to Edit Mode to edit text")
            return
        if self.selected_span and self.selected_page == self.current_page:
            self._start_inline_edit()

    def _start_inline_edit(self):
        if not self.selected_span:
            return
        self._cancel_inline_edit()
        span = self.selected_span
        r = span["bbox"] * self.zoom
        # Format bar above the text
        self._fmt_frame = tk.Frame(self.canvas, bg="#2a2a3c", padx=4, pady=2)
        tk.Label(self._fmt_frame, text="Size:", bg="#2a2a3c", fg="#e8e8f0", font=("Segoe UI", 8)).pack(side=tk.LEFT)
        self._fmt_size = tk.StringVar(value=str(int(round(float(span["size"])))))
        tk.Button(self._fmt_frame, text="−", command=lambda: self._bump_fmt_size(-1),
                  bg="#34344a", fg="#e8e8f0", relief=tk.FLAT, font=("Segoe UI", 9, "bold"),
                  width=2, padx=2).pack(side=tk.LEFT)
        size_entry = tk.Entry(self._fmt_frame, textvariable=self._fmt_size, width=4,
                              bg="#34344a", fg="#e8e8f0", relief=tk.FLAT, font=("Segoe UI", 9),
                              justify=tk.CENTER)
        size_entry.pack(side=tk.LEFT, padx=2)
        tk.Button(self._fmt_frame, text="+", command=lambda: self._bump_fmt_size(1),
                  bg="#34344a", fg="#e8e8f0", relief=tk.FLAT, font=("Segoe UI", 9, "bold"),
                  width=2, padx=2).pack(side=tk.LEFT)
        self._fmt_bold = tk.BooleanVar(value=bool(span.get("flags", 0) & 16))
        self._fmt_italic = tk.BooleanVar(value=bool(span.get("flags", 0) & 2))
        tk.Checkbutton(self._fmt_frame, text="B", variable=self._fmt_bold, bg="#2a2a3c", fg="#e8e8f0",
                       selectcolor="#7c5cff", font=("Segoe UI", 9, "bold"),
                       activebackground="#2a2a3c").pack(side=tk.LEFT, padx=2)
        tk.Checkbutton(self._fmt_frame, text="I", variable=self._fmt_italic, bg="#2a2a3c", fg="#e8e8f0",
                       selectcolor="#7c5cff", font=("Segoe UI", 9, "italic"),
                       activebackground="#2a2a3c").pack(side=tk.LEFT, padx=2)
        tk.Button(self._fmt_frame, text="বাংলা", command=self._open_bangla_kb,
                  bg="#0ea5e9", fg="white", relief=tk.FLAT, font=("Segoe UI", 8), padx=6).pack(side=tk.LEFT, padx=4)
        tk.Button(self._fmt_frame, text="Apply", command=self._finish_inline_edit,
                  bg="#7c5cff", fg="white", relief=tk.FLAT, font=("Segoe UI", 8), padx=6).pack(side=tk.LEFT, padx=4)
        tk.Button(self._fmt_frame, text="X", command=self._cancel_inline_edit,
                  bg="#34344a", fg="#e8e8f0", relief=tk.FLAT, font=("Segoe UI", 8), padx=4).pack(side=tk.LEFT)
        self.canvas.create_window(r.x0, max(0, r.y0 - 28), window=self._fmt_frame, anchor=tk.NW, tags="editentry")
        self._edit_entry = tk.Entry(
            self.canvas,
            font=("Segoe UI", max(9, int(span["size"] * self.zoom * 0.75))),
            bg="#fffde7", fg="#111", relief=tk.SOLID, bd=1, insertbackground="#111",
        )
        self._edit_entry.insert(0, span["text"])
        self._edit_entry.select_range(0, tk.END)
        self._edit_span = span
        self.canvas.create_window(r.x0, r.y0, window=self._edit_entry, anchor=tk.NW,
                                  width=max(r.width, 120), height=max(r.height + 4, 22), tags="editentry")
        self._edit_entry.focus_set()
        self._edit_entry.bind("<Return>", lambda e: self._finish_inline_edit())
        self._edit_entry.bind("<Escape>", lambda e: self._cancel_inline_edit())
        self.status.config(text="Edit text + set Size / Bold / Italic — Apply or Enter")

    def _bump_fmt_size(self, delta):
        try:
            cur = float(self._fmt_size.get())
        except Exception:
            cur = 12.0
        cur = max(4.0, min(96.0, cur + delta))
        self._fmt_size.set(str(int(cur) if cur == int(cur) else cur))

    def _open_bangla_kb(self):
        """Open on-screen Bangla keyboard targeting the inline edit Entry."""
        if not self._edit_entry:
            messagebox.showinfo("Edit first", "Double-click text to edit, then open বাংলা keyboard.")
            return

        def insert_char(ch):
            if not self._edit_entry:
                return
            if ch == "BACK":
                try:
                    self._edit_entry.delete(len(self._edit_entry.get()) - 1, tk.END)
                except Exception:
                    pass
                return
            self._edit_entry.insert(tk.INSERT, ch)
            self._edit_entry.focus_set()

        BanglaKeyboard(self, on_char=insert_char, on_enter=self._finish_inline_edit)

    def _finish_inline_edit(self):
        if not self._edit_entry or not self._edit_span:
            self._cancel_inline_edit()
            return
        new_text = self._edit_entry.get()
        span = self._edit_span
        page = self.selected_page
        try:
            fs = float(self._fmt_size.get())
            fs = max(4.0, min(96.0, fs))
        except Exception:
            fs = float(span["size"])
        bold = bool(self._fmt_bold.get()) if hasattr(self, "_fmt_bold") else False
        italic = bool(self._fmt_italic.get()) if hasattr(self, "_fmt_italic") else False
        self._cancel_inline_edit()
        if not new_text.strip():
            return
        ok, msg = self.pdf.replace_span(page, span, new_text, fontsize=fs, bold=bold, italic=italic)
        self.selected_span = None
        if ok:
            self.render_page()
            self.status.config(text=msg)
        else:
            messagebox.showerror("Edit failed", msg)

    def _cancel_inline_edit(self):
        if self._edit_entry:
            try:
                self._edit_entry.destroy()
            except Exception:
                pass
        if hasattr(self, "_fmt_frame") and self._fmt_frame:
            try:
                self._fmt_frame.destroy()
            except Exception:
                pass
            self._fmt_frame = None
        self._edit_entry = None
        self._edit_span = None
        try:
            self.canvas.delete("editentry")
        except Exception:
            pass

    def _on_canvas_drag(self, event):
        cx = self.canvas.canvasx(event.x)
        cy = self.canvas.canvasy(event.y)
        pt = self._canvas_to_pdf(cx, cy)
        # Move selected text
        if self._move_active and self.selected_span and self._move_start is not None:
            if not (self.screenshot_mode or self.fill_mode or self.copy_sign_mode or self.placing_signature):
                self._move_delta = (pt.x - self._move_start.x, pt.y - self._move_start.y)
                self.render_page()
                return
        if not self._drag_start:
            return
        if self.screenshot_mode:
            self.screenshot_rect = fitz.Rect(self._drag_start, pt)
            self.screenshot_rect.normalize()
            self.render_page()
            return
        if self.copy_sign_mode:
            self.copy_sign_rect = fitz.Rect(self._drag_start, pt)
            self.copy_sign_rect.normalize()
            self.render_page()
            return
        if self.link_mode:
            self.link_rect = fitz.Rect(self._drag_start, pt)
            self.link_rect.normalize()
            self.render_page()
            return
        if self.fill_mode:
            self.fill_rect = fitz.Rect(self._drag_start, pt)
            self.fill_rect.normalize()
            self.render_page()
            return
        if self.placing_signature:
            self.sig_rect = fitz.Rect(self._drag_start, pt)
            self.sig_rect.normalize()
            self.render_page()

    def _on_canvas_release(self, event):
        # Finish text move
        if self._move_active and self.selected_span and self._move_start is not None:
            dx, dy = self._move_delta
            self._move_active = False
            self._move_start = None
            self._drag_start = None
            try:
                self.canvas.config(cursor="")
            except Exception:
                pass
            if abs(dx) >= 1.0 or abs(dy) >= 1.0:
                span = self.selected_span
                page = self.selected_page
                ok, msg = self.pdf.move_span(page, span, dx, dy)
                self._move_delta = (0.0, 0.0)
                self.selected_span = None
                if ok:
                    self.render_page()
                    self.status.config(text="Text moved — click to select again")
                else:
                    self.status.config(text=f"Move: {msg}")
                    self.render_page()
            else:
                self._move_delta = (0.0, 0.0)
                self.render_page()
            return
        if self.screenshot_mode and self.screenshot_rect:
            self._screenshot_selected_area()
            return
        if self.copy_sign_mode and self.copy_sign_rect:
            self._capture_sign_region()
            return
        if self.link_mode and self.link_rect:
            self._finish_add_hyperlink()
            return
        if self.fill_mode and self.fill_rect:
            if self.fill_rect.width < 5 or self.fill_rect.height < 5:
                self.status.config(text="Box too small")
                return
            self._apply_fill_box()
            return
        if self.placing_signature and self.sig_rect and self.signature_img:
            if self.sig_rect.width < 10 or self.sig_rect.height < 5:
                self.status.config(text="Signature area too small")
                return
            ok = self.pdf.insert_signature_image(self.current_page, self.signature_img, self.sig_rect)
            if ok:
                self.status.config(text="Signature placed")
                self.placing_signature = False
                self.signature_img = None
                self.sig_rect = None
                self._drag_start = None
                self.render_page()
            else:
                messagebox.showerror("Error", "Could not insert signature")

    def show_search(self):
        self.search_frame.pack(side=tk.TOP, fill=tk.X)
        self.search_entry.focus_set()
        self.search_entry.select_range(0, tk.END)

    def hide_search(self):
        self.search_frame.pack_forget()
        self.highlight_rect = None
        self.render_page()

    def do_search(self):
        q = self.search_var.get().strip()
        if not q or not self.pdf.doc:
            return
        self.search_results = self.pdf.search(q)
        self.search_index = -1
        if not self.search_results:
            self.search_status.config(text="No matches")
            self.highlight_rect = None
            self.render_page()
        else:
            self.search_status.config(text=f"{len(self.search_results)} found")
            self.find_next()

    def find_next(self):
        if not self.search_results:
            self.do_search()
            return
        self.search_index = (self.search_index + 1) % len(self.search_results)
        self._goto_search_hit()

    def find_prev(self):
        if not self.search_results:
            return
        self.search_index = (self.search_index - 1) % len(self.search_results)
        self._goto_search_hit()

    def _goto_search_hit(self):
        page_idx, rect = self.search_results[self.search_index]
        self.current_page = page_idx
        self.highlight_rect = rect
        self._update_page_ui()
        self.render_page()
        self.search_status.config(text=f"{self.search_index + 1} / {len(self.search_results)}")

    def start_signature(self):
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a file first.")
            return
        def on_done(img):
            if img is None:
                return
            self.signature_img = img
            self.placing_signature = True
            self.sig_rect = None
            self._drag_start = None
            self.status.config(text="Drag a rectangle on the page to place the signature")
        SignatureDialog(self, on_done)




    def _require_edit_mode(self, action="edit"):
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a file first.", parent=self)
            return False
        if not self.edit_mode:
            if messagebox.askyesno(
                "Edit Mode",
                f"File is in View Mode.\n\nSwitch to Edit Mode to {action}?",
                parent=self,
            ):
                self.set_edit_mode(True)
                return True
            return False
        return True

    def _on_escape(self, event=None):
        if self.link_edit_mode:
            self.link_edit_mode = False
            try:
                self.canvas.config(cursor="")
            except Exception:
                pass
            self.render_page()
            self.status.config(text="Link edit cancelled")
            return
        if self.link_mode or self.fill_mode or self.screenshot_mode or self.copy_sign_mode:
            self._cancel_ops()
            self.render_page()
            self.status.config(text="Tool cancelled")
            return
        self._cancel_inline_edit()

    def toggle_edit_mode(self):
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a file first.", parent=self)
            return
        self.set_edit_mode(not self.edit_mode)

    def set_edit_mode(self, enabled):
        self.edit_mode = bool(enabled)
        self.link_edit_mode = False
        if not self.edit_mode:
            self._cancel_ops()
            self.selected_span = None
            self._move_active = False
        try:
            if self.edit_mode:
                self.btn_edit_mode.config(text="Edit Mode", bg="#22c55e", fg="white")
                self.status.config(text="EDIT mode — text, tools, links editable")
            else:
                self.btn_edit_mode.config(text="View Mode", bg="#0ea5e9", fg="white")
                self.status.config(text="VIEW mode — click links to open in browser · press Edit Mode to change PDF")
        except Exception:
            pass
        self.render_page()

    def _open_link_at(self, page_idx, pt):
        """If point hits a URI link, open it in the default browser. Returns True if handled."""
        try:
            links = self.pdf.get_page_links(page_idx)
        except Exception:
            return False
        for lk in links:
            try:
                r = fitz.Rect(lk.get("from"))
                if not r.contains(pt):
                    continue
                uri = (lk.get("uri") or "").strip()
                if uri:
                    if not (uri.startswith("http://") or uri.startswith("https://") or uri.startswith("mailto:")):
                        if "@" in uri and " " not in uri:
                            uri = "mailto:" + uri
                        else:
                            uri = "https://" + uri
                    webbrowser.open(uri)
                    self.status.config(text=f"Opened: {uri[:80]}")
                    return True
                # internal go-to page
                if lk.get("page") is not None:
                    dest = int(lk.get("page"))
                    if 0 <= dest < self.pdf.page_count():
                        self.current_page = dest
                        self._update_page_ui()
                        self.render_page()
                        self.scroll_to_page(dest)
                        self.status.config(text=f"Jumped to page {dest + 1}")
                        return True
            except Exception:
                continue
        return False

    def manage_hyperlinks(self):
        """One window: click a link in the list → edit URL here → Save URL."""
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a PDF first.", parent=self)
            return
        if not self._require_edit_mode("edit hyperlinks"):
            return

        page_idx = self.current_page
        self.link_edit_mode = False  # edit only inside this window, not separate popups
        try:
            self.canvas.config(cursor="")
        except Exception:
            pass
        self.render_page()

        win = tk.Toplevel(self)
        win.title(f"Hyperlinks — Page {page_idx + 1}")
        win.configure(bg=COLORS["surface"])
        win.geometry("680x480")
        win.minsize(600, 420)
        win.transient(self)
        win.resizable(True, True)

        # ---- layout: top instructions, middle list, bottom editor+buttons (always visible)
        top = tk.Frame(win, bg=COLORS["surface"])
        top.pack(side=tk.TOP, fill=tk.X, padx=16, pady=(14, 6))
        tk.Label(
            top, text="Hyperlinks on this page",
            bg=COLORS["surface"], fg=COLORS["text"], font=("Segoe UI", 14, "bold"),
        ).pack(anchor=tk.W)
        tk.Label(
            top,
            text="Click a link in the list → URL appears below → change it → press Save URL",
            bg=COLORS["surface"], fg=COLORS["text_dim"], font=("Segoe UI", 9),
        ).pack(anchor=tk.W, pady=(4, 0))

        bottom = tk.Frame(win, bg=COLORS["surface"])
        bottom.pack(side=tk.BOTTOM, fill=tk.X, padx=16, pady=14)

        mid = tk.Frame(win, bg=COLORS["surface"])
        mid.pack(side=tk.TOP, fill=tk.BOTH, expand=True, padx=16, pady=8)

        lb = tk.Listbox(
            mid,
            bg=COLORS["surface2"],
            fg=COLORS["text"],
            font=("Segoe UI", 11),
            selectbackground=COLORS["accent"],
            selectforeground="white",
            relief=tk.FLAT,
            activestyle="none",
            highlightthickness=0,
            exportselection=False,
        )
        sb = tk.Scrollbar(mid, orient=tk.VERTICAL, command=lb.yview)
        lb.configure(yscrollcommand=sb.set)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        lb.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        items = []

        tk.Label(
            bottom, text="Edit URL (same window):",
            bg=COLORS["surface"], fg=COLORS["text"], font=("Segoe UI", 9, "bold"),
        ).pack(anchor=tk.W)
        url_var = tk.StringVar()
        url_entry = tk.Entry(
            bottom,
            textvariable=url_var,
            font=("Segoe UI", 12),
            bg="#1e1e2e",
            fg="#ffffff",
            insertbackground="#ffffff",
            relief=tk.SOLID,
            bd=1,
        )
        url_entry.pack(fill=tk.X, pady=(6, 10), ipady=10)

        btn_row = tk.Frame(bottom, bg=COLORS["surface"])
        btn_row.pack(fill=tk.X)

        def refresh(select=0):
            nonlocal items
            lb.delete(0, tk.END)
            try:
                items = list(self.pdf.get_page_links(page_idx))
            except Exception:
                items = []
            if not items:
                lb.insert(tk.END, "(No links on this page — use Add New)")
                url_var.set("")
                return
            for i, lk in enumerate(items):
                uri = (lk.get("uri") or "").strip()
                if uri:
                    lb.insert(tk.END, f"{i + 1}.  {uri}")
                elif lk.get("page") is not None:
                    lb.insert(tk.END, f"{i + 1}.  Go to page {int(lk.get('page')) + 1}")
                else:
                    lb.insert(tk.END, f"{i + 1}.  (link)")
            select = max(0, min(int(select), len(items) - 1))
            lb.selection_clear(0, tk.END)
            lb.selection_set(select)
            lb.activate(select)
            lb.see(select)
            on_select()

        def on_select(event=None):
            sel = lb.curselection()
            if not sel or not items:
                return
            i = int(sel[0])
            if i < 0 or i >= len(items):
                return
            uri = (items[i].get("uri") or "").strip()
            url_var.set(uri)
            url_entry.focus_set()
            try:
                url_entry.selection_range(0, tk.END)
            except Exception:
                pass
            # highlight on PDF page
            try:
                self.highlight_rect = fitz.Rect(items[i].get("from"))
                self.current_page = page_idx
                self.render_page()
            except Exception:
                pass

        def current():
            sel = lb.curselection()
            if not sel or not items:
                return None, -1
            i = int(sel[0])
            if i < 0 or i >= len(items):
                return None, -1
            return items[i], i

        def do_save():
            lk, i = current()
            if lk is None:
                messagebox.showinfo("Select a link", "Click a link in the list first.", parent=win)
                return
            new_uri = url_var.get().strip()
            if not new_uri:
                messagebox.showinfo("URL empty", "Type the new URL in the box below the list.", parent=win)
                return
            ok, msg = self.pdf.update_link_uri(page_idx, lk, new_uri)
            if ok:
                refresh(select=i)
                self.status.config(text=msg + " — File → Save to write PDF")
                try:
                    win.title(f"Hyperlinks — Page {page_idx + 1}  ✓ saved")
                except Exception:
                    pass
            else:
                messagebox.showerror("Could not update", msg, parent=win)

        def do_delete():
            lk, i = current()
            if lk is None:
                messagebox.showinfo("Select a link", "Click a link in the list first.", parent=win)
                return
            if not messagebox.askyesno("Delete", "Remove this hyperlink from the page?", parent=win):
                return
            ok, msg = self.pdf.delete_link(page_idx, lk)
            if ok:
                self.highlight_rect = None
                self.render_page()
                refresh(select=0)
                self.status.config(text=msg)
            else:
                messagebox.showerror("Could not delete", msg, parent=win)

        def do_add():
            win.destroy()
            self.start_add_hyperlink()

        def do_close():
            self.highlight_rect = None
            try:
                self.render_page()
            except Exception:
                pass
            win.destroy()

        tk.Button(
            btn_row, text="Save URL", command=do_save,
            bg=COLORS["accent"], fg="white", relief=tk.FLAT,
            font=("Segoe UI", 10, "bold"), padx=18, pady=9, cursor="hand2",
        ).pack(side=tk.LEFT)
        tk.Button(
            btn_row, text="Delete", command=do_delete,
            bg="#ef4444", fg="white", relief=tk.FLAT,
            font=("Segoe UI", 9), padx=12, pady=9, cursor="hand2",
        ).pack(side=tk.LEFT, padx=(8, 0))
        tk.Button(
            btn_row, text="Add New Link…", command=do_add,
            bg=COLORS["surface2"], fg=COLORS["text"], relief=tk.FLAT,
            font=("Segoe UI", 9), padx=12, pady=9, cursor="hand2",
        ).pack(side=tk.LEFT, padx=(8, 0))
        tk.Button(
            btn_row, text="Close", command=do_close,
            bg=COLORS["surface2"], fg=COLORS["text_dim"], relief=tk.FLAT,
            font=("Segoe UI", 9), padx=12, pady=9, cursor="hand2",
        ).pack(side=tk.RIGHT)

        lb.bind("<<ListboxSelect>>", on_select)
        lb.bind("<ButtonRelease-1>", on_select)
        win.bind("<Return>", lambda e: do_save())
        win.protocol("WM_DELETE_WINDOW", do_close)

        refresh(0)
        if items:
            on_select()
        # ensure widgets laid out
        win.update_idletasks()

    def _edit_link_at(self, page_idx, pt):
        """Click link on page while link_edit_mode: open same URL editor."""
        try:
            links = self.pdf.get_page_links(page_idx)
        except Exception:
            return False
        hit = None
        for lk in links:
            try:
                if fitz.Rect(lk.get("from")).contains(pt):
                    hit = lk
                    break
            except Exception:
                continue
        if not hit:
            return False
        # Reuse small dialog
        cur = (hit.get("uri") or "").strip()
        win = tk.Toplevel(self)
        win.title("Edit Hyperlink")
        win.configure(bg=COLORS["surface"])
        win.minsize(420, 160)
        win.geometry("500x200")
        win.transient(self)
        tk.Label(
            win, text="URL (text on page will not change):",
            bg=COLORS["surface"], fg=COLORS["text"], font=("Segoe UI", 10),
        ).pack(anchor=tk.W, padx=16, pady=(16, 6))
        var = tk.StringVar(value=cur)
        ent = tk.Entry(
            win, textvariable=var, font=("Segoe UI", 12),
            bg="#1e1e2e", fg="#f8f8ff", insertbackground="#f8f8ff",
            relief=tk.SOLID, bd=1,
        )
        ent.pack(fill=tk.X, padx=16, ipady=10)
        ent.focus_set()
        ent.selection_range(0, tk.END)
        bf = tk.Frame(win, bg=COLORS["surface"])
        bf.pack(fill=tk.X, padx=16, pady=16)

        def save():
            new_uri = var.get().strip()
            if not new_uri:
                return
            ok, msg = self.pdf.update_link_uri(page_idx, hit, new_uri)
            if ok:
                self.render_page()
                self.status.config(text=msg + " — File → Save")
                win.destroy()
            else:
                messagebox.showerror("Failed", msg, parent=win)

        tk.Button(bf, text="Save Link", command=save, bg=COLORS["accent"], fg="white",
                  relief=tk.FLAT, padx=14, pady=8, font=("Segoe UI", 10, "bold"), cursor="hand2").pack(side=tk.LEFT)
        tk.Button(bf, text="Cancel", command=win.destroy, bg=COLORS["surface2"], fg=COLORS["text"],
                  relief=tk.FLAT, padx=12, pady=8, font=("Segoe UI", 9), cursor="hand2").pack(side=tk.RIGHT)
        win.bind("<Return>", lambda e: save())
        return True

    def add_link_to_selected_text(self):
        """Make the currently selected text span a clickable hyperlink."""
        if not self._require_edit_mode("add link on text"):
            return
        if not self.selected_span or self.selected_page < 0:
            messagebox.showinfo(
                "Select text",
                "First click the text (e.g. FORM) so it is selected (green box),\n"
                "then use Tools → Add Link to Selected Text…",
                parent=self,
            )
            return
        span = self.selected_span
        page_idx = self.selected_page
        sample = (span.get("text") or "").strip() or "text"
        uri = simpledialog.askstring(
            "Link for selected text",
            f'Text: "{sample[:60]}"\n\nEnter URL (https://...):',
            parent=self,
        )
        if not uri:
            return
        rect = fitz.Rect(span["bbox"])
        # slight pad so click target is comfortable
        rect = fitz.Rect(rect.x0 - 1, rect.y0 - 1, rect.x1 + 1, rect.y1 + 1)
        ok, msg = self.pdf.add_uri_link(page_idx, rect, uri)
        if ok:
            self.render_page()
            self.status.config(text=f'Link on "{sample[:30]}": {msg}')
        else:
            messagebox.showerror("Link failed", msg, parent=self)

    def start_add_hyperlink(self):
        if not self._require_edit_mode("add hyperlinks"):
            return
        self._cancel_ops()
        self.link_edit_mode = False
        self.link_mode = True
        self.link_rect = None
        try:
            self.canvas.config(cursor="crosshair")
        except Exception:
            pass
        self.status.config(text="Add hyperlink: drag a rectangle on the page, then enter URL")

    def _finish_add_hyperlink(self):
        if not self.link_rect or self.link_rect.width < 3:
            self.link_mode = False
            self.link_rect = None
            try:
                self.canvas.config(cursor="")
            except Exception:
                pass
            self.status.config(text="Link area too small")
            return
        uri = simpledialog.askstring(
            "Hyperlink URL",
            "Enter URL (example: https://example.com):",
            parent=self,
        )
        rect = fitz.Rect(self.link_rect)
        self.link_mode = False
        self.link_rect = None
        self.link_edit_mode = False
        try:
            self.canvas.config(cursor="")
        except Exception:
            pass
        if not uri:
            self.render_page()
            self.status.config(text="Add link cancelled")
            return
        ok, msg = self.pdf.add_uri_link(self.current_page, rect, uri)
        self.render_page()
        if ok:
            self.status.config(text=msg)
        else:
            messagebox.showerror("Link failed", msg, parent=self)

    def take_screenshot(self):
        if not self.pdf.doc:
            messagebox.showinfo("No page", "Open a file first.")
            return
        # Choice dialog
        win = tk.Toplevel(self)
        win.title("Screenshot")
        win.configure(bg=COLORS["surface"])
        win.transient(self)
        win.grab_set()
        win.resizable(False, False)
        tk.Label(win, text="Choose screenshot type", bg=COLORS["surface"], fg=COLORS["text"],
                 font=("Segoe UI", 12, "bold")).pack(padx=20, pady=(16, 8))
        bf = tk.Frame(win, bg=COLORS["surface"])
        bf.pack(padx=20, pady=(4, 16))

        def do_page():
            win.destroy()
            self._screenshot_full_page()

        def do_area():
            win.destroy()
            self.screenshot_mode = True
            self.fill_mode = False
            self.copy_sign_mode = False
            self.placing_signature = False
            self.pick_color_mode = False
            self.screenshot_rect = None
            self._drag_start = None
            try:
                self.canvas.config(cursor="crosshair")
            except Exception:
                pass
            self.status.config(text="Screenshot area: drag a rectangle on the page")

        tk.Button(bf, text="📄  Current Page", command=do_page,
                  bg=COLORS["accent"], fg="white", relief=tk.FLAT, padx=18, pady=10,
                  font=("Segoe UI", 10), cursor="hand2").pack(side=tk.LEFT, padx=6)
        tk.Button(bf, text="⬚  Select Area", command=do_area,
                  bg=COLORS["surface2"], fg=COLORS["text"], relief=tk.FLAT, padx=18, pady=10,
                  font=("Segoe UI", 10), cursor="hand2").pack(side=tk.LEFT, padx=6)
        tk.Button(win, text="Cancel", command=win.destroy,
                  bg=COLORS["surface2"], fg=COLORS["text_dim"], relief=tk.FLAT, padx=12, pady=6,
                  font=("Segoe UI", 9), cursor="hand2").pack(pady=(0, 14))
        win.update_idletasks()
        x = self.winfo_rootx() + (self.winfo_width() - win.winfo_width()) // 2
        y = self.winfo_rooty() + (self.winfo_height() - win.winfo_height()) // 2
        win.geometry(f"+{x}+{y}")

    def _screenshot_full_page(self):
        try:
            page = self.pdf.get_page(self.current_page)
            if not page:
                return
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            ScreenshotPopup(self, img)
        except Exception as e:
            messagebox.showerror("Screenshot failed", str(e))

    def _screenshot_selected_area(self):
        rect = self.screenshot_rect
        self.screenshot_mode = False
        self.screenshot_rect = None
        self._drag_start = None
        try:
            self.canvas.config(cursor="")
        except Exception:
            pass
        if not rect or rect.width < 5 or rect.height < 5:
            self.status.config(text="Screenshot area too small")
            self.render_page()
            return
        try:
            page = self.pdf.get_page(self.current_page)
            if not page:
                return
            mat = fitz.Matrix(2.5, 2.5)
            pix = page.get_pixmap(matrix=mat, clip=rect, alpha=False)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            ScreenshotPopup(self, img)
            self.status.config(text="Area screenshot ready")
            self.render_page()
        except Exception as e:
            messagebox.showerror("Screenshot failed", str(e))

    def _cancel_ops(self):
        self._cancel_inline_edit()
        self.placing_signature = False
        self.signature_img = None
        self.sig_rect = None
        self.fill_mode = False
        self.pick_color_mode = False
        self.copy_sign_mode = False
        self.screenshot_mode = False
        self.link_mode = False
        self.link_rect = None
        self.link_edit_mode = False
        self.placing_symbol = False
        self.symbol_char = None
        self.fill_rect = None
        self.copy_sign_rect = None
        self.screenshot_rect = None
        self._drag_start = None
        try:
            self.canvas.config(cursor="")
        except Exception:
            pass
        self.hide_search()
        self.status.config(text="Cancelled")


    def start_pick_color(self):
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a file first.")
            return
        self.pick_color_mode = True
        self.fill_mode = False
        self.placing_signature = False
        self.status.config(text="Click anywhere on the page to pick that color")

    def _sample_color_at(self, canvas_x, canvas_y):
        """Sample RGB from current rendered page pixmap."""
        try:
            page = self.pdf.get_page(self.current_page)
            if not page:
                return
            mat = fitz.Matrix(self.zoom, self.zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            x = int(canvas_x)
            y = int(canvas_y)
            if 0 <= x < pix.width and 0 <= y < pix.height:
                # pix.samples is RGB
                i = (y * pix.width + x) * 3
                r = pix.samples[i] / 255.0
                g = pix.samples[i + 1] / 255.0
                b = pix.samples[i + 2] / 255.0
                self.fill_color = (r, g, b)
                self.fill_color_manual = True
                self.pick_color_mode = False
                hexc = "#%02x%02x%02x" % (int(r*255), int(g*255), int(b*255))
                self.status.config(text=f"Fill color picked {hexc} — use Fill Box")
                messagebox.showinfo("Color Picked", f"Fill color: {hexc}\n\nNow use Fill Box and drag a rectangle.")
            else:
                self.status.config(text="Click inside the page")
        except Exception as e:
            self.status.config(text=f"Pick failed: {e}")
            self.pick_color_mode = False

    def start_fill_box(self):
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a file first.")
            return
        self.fill_mode = True
        self.pick_color_mode = False
        self.placing_signature = False
        self.copy_sign_mode = False
        self.fill_rect = None
        self._drag_start = None
        if self.fill_color_manual and self.fill_color:
            c = self.fill_color
            hexc = "#%02x%02x%02x" % (int(c[0]*255), int(c[1]*255), int(c[2]*255))
            self.status.config(text=f"Fill Box (manual {hexc}): drag rectangle")
        else:
            self.status.config(text="Fill Box (auto match PDF background): drag rectangle")

    def reset_fill_color(self):
        self.fill_color = None
        self.fill_color_manual = False
        self.status.config(text="Fill color reset — will auto-match PDF background")

    def choose_text_color(self):
        initial = (
            int(self.text_color[0] * 255),
            int(self.text_color[1] * 255),
            int(self.text_color[2] * 255),
        )
        result = colorchooser.askcolor(color=initial, title="Choose text color")
        if result and result[0]:
            r, g, b = result[0]
            self.text_color = (r / 255.0, g / 255.0, b / 255.0)
            self.status.config(text=f"Text color set to #{int(r):02x}{int(g):02x}{int(b):02x}")

    def _apply_fill_box(self):
        rect = self.fill_rect
        if not rect:
            return
        # Keep a copy — doc changes must not lose coordinates
        saved_rect = fitz.Rect(rect)
        # Auto-sample background unless user manually picked a fill color
        if self.fill_color_manual and self.fill_color is not None:
            use_color = self.fill_color
        else:
            use_color = self.pdf.sample_rect_background(self.current_page, saved_rect)
        ok = self.pdf.fill_rectangle(self.current_page, saved_rect, use_color)
        self.fill_mode = False
        self.fill_rect = None
        self._drag_start = None
        if not ok:
            messagebox.showerror("Error", "Could not apply fill box")
            self.render_page()
            return
        self.render_page()
        if messagebox.askyesno("Write text?", "Area covered.\n\nDo you want to type new text on this area?"):
            text = self._ask_text_with_bangla("Enter text (English or use বাংলা keyboard):")
            if text and text.strip():
                ok2 = self.pdf.insert_text_in_rect(
                    self.current_page, saved_rect, text.strip(),
                    fontsize=None, color=self.text_color,
                )
                self.render_page()
                if ok2:
                    self.status.config(text="Fill box + new text applied")
                else:
                    messagebox.showerror("Text failed", "Could not insert text. Try a larger box.")
                    self.status.config(text="Fill applied but text insert failed")
            else:
                self.status.config(text="Fill box applied (text hidden)")
        else:
            self.status.config(text="Fill box applied (text hidden)")

    def copy_signature_from_pdf(self):
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a file first.")
            return
        images = self.pdf.extract_page_images(self.current_page)
        if not images:
            messagebox.showinfo("No images", "No embedded images found on this page.\nSignatures drawn as vector may not appear.")
            return
        # Simple chooser dialog
        win = tk.Toplevel(self)
        win.title("Copy Signature from Page")
        win.configure(bg=COLORS["surface"])
        win.transient(self)
        win.grab_set()
        tk.Label(win, text="Select an image to use as signature:", bg=COLORS["surface"],
                 fg=COLORS["text"], font=("Segoe UI", 10)).pack(padx=12, pady=8)
        frame = tk.Frame(win, bg=COLORS["surface"])
        frame.pack(padx=12, pady=4)
        self._sig_choice_imgs = []
        for idx, (xref, img) in enumerate(images[:12]):
            thumb = img.copy()
            thumb.thumbnail((120, 60))
            photo = ImageTk.PhotoImage(thumb)
            self._sig_choice_imgs.append(photo)
            def make_cmd(im=img):
                def cmd():
                    self.signature_img = im.convert("RGBA") if im.mode != "RGBA" else im
                    self.placing_signature = True
                    self.sig_rect = None
                    self._drag_start = None
                    self.status.config(text="Signature copied — drag a rectangle to place it")
                    win.destroy()
                return cmd
            btn = tk.Button(frame, image=photo, command=make_cmd(), bg=COLORS["surface2"], relief=tk.FLAT)
            btn.grid(row=idx // 4, column=idx % 4, padx=4, pady=4)
        tk.Button(win, text="Cancel", command=win.destroy, bg=COLORS["surface2"], fg=COLORS["text"],
                  relief=tk.FLAT, padx=12, pady=6).pack(pady=10)
        win.geometry("560x280")

    def show_default_viewer_help(self):
        messagebox.showinfo(
            "Set as Default PDF Viewer",
            "To open PDFs with One PDF Editor by default on Windows:\n\n"
            "1. Right-click any PDF file\n"
            "2. Choose Open with → Choose another app\n"
            "3. Browse and select OnePDFEditor.exe\n"
            "4. Check Always use this app to open .pdf files\n"
            "5. Click OK\n\n"
            "Or: Settings → Apps → Default apps → PDF → One PDF Editor\n\n"
            "This must be done once on your PC after you have the .exe.",
        )


    def ocr_current_page(self):
        """Try to OCR current page so image text becomes selectable/editable."""
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a file first.")
            return
        page = self.pdf.get_page(self.current_page)
        if not page:
            return
        try:
            self.status.config(text="Running OCR… please wait")
            self.update_idletasks()
            self.pdf.save_state()
            # PyMuPDF OCR (requires Tesseract installed on the PC)
            tp = page.get_textpage_ocr(dpi=200, full=True)
            text = page.get_text("text", textpage=tp)
            if not text.strip():
                messagebox.showinfo(
                    "OCR result",
                    "No text detected.\n\n"
                    "For OCR you need Tesseract installed on Windows:\n"
                    "https://github.com/UB-Mannheim/tesseract/wiki\n\n"
                    "Without Tesseract, you can still:\n"
                    "• Cover areas with Fill Box\n"
                    "• Type new text on top\n"
                    "• Add signatures",
                )
                self.status.config(text="OCR: no text found (is Tesseract installed?)")
                return
            # Re-insert detected text as real PDF text (simple full-page overlay approach)
            # Better: user can now search; for editing, spans from OCR textpage help
            messagebox.showinfo(
                "OCR done",
                f"Detected text ({len(text)} chars).\n\n"
                "You can now Search the text.\n"
                "For precise edit: use Fill Box to cover old glyphs, then type new text.\n\n"
                "Preview of detected text:\n" + text[:300],
            )
            self.status.config(text="OCR finished — search works on detected text")
            self.render_page()
        except Exception as e:
            messagebox.showinfo(
                "OCR not available",
                "Could not run OCR.\n\n"
                "Install Tesseract for Windows, then restart the app.\n"
                "Download: https://github.com/UB-Mannheim/tesseract/wiki\n\n"
                f"Error: {e}\n\n"
                "Without OCR you can still hide text (Fill Box) and write new text.",
            )
            self.status.config(text="OCR unavailable — install Tesseract for image text edit")


    def start_copy_sign_region(self):
        """Plus-cursor: drag a rectangle; capture that area as signature."""
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a file first.")
            return
        self.copy_sign_mode = True
        self.fill_mode = False
        self.pick_color_mode = False
        self.placing_signature = False
        self.copy_sign_rect = None
        self._drag_start = None
        try:
            self.canvas.config(cursor="crosshair")
        except Exception:
            pass
        self.status.config(text="Copy Sign: drag a rectangle over the signature area")

    def _capture_sign_region(self):
        rect = self.copy_sign_rect
        if not rect or rect.width < 5 or rect.height < 5:
            self.status.config(text="Area too small")
            return
        try:
            page = self.pdf.get_page(self.current_page)
            # high-res capture of region
            mat = fitz.Matrix(2.5, 2.5)
            pix = page.get_pixmap(matrix=mat, clip=rect, alpha=True)
            mode = "RGBA" if pix.alpha else "RGB"
            img = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
            if img.mode != "RGBA":
                img = img.convert("RGBA")
            self.signature_img = img
            self.copy_sign_mode = False
            self.copy_sign_rect = None
            self._drag_start = None
            self.placing_signature = True
            try:
                self.canvas.config(cursor="")
            except Exception:
                pass
            self.status.config(text="Signature captured — now drag where you want to place it")
            self.render_page()
        except Exception as e:
            messagebox.showerror("Copy failed", str(e))
            self.copy_sign_mode = False
            try:
                self.canvas.config(cursor="")
            except Exception:
                pass


    def _show_dashboard(self):
        """Home screen: near-realtime video frame playback + transparent text."""
        try:
            self.page_count_lbl.config(text="/ -")
            self.page_var.set("-")
        except Exception:
            pass
        try:
            self.status.config(text="")
        except Exception:
            pass

        # Load all video frames once (12 fps sequence)
        if not getattr(self, "_dash_photos", None):
            self._dash_photos = []
            roots = [resource_path("assets", "dashboard")]
            if getattr(sys, "frozen", False):
                roots.extend([
                    Path(sys.executable).parent / "assets" / "dashboard",
                    Path(sys.executable).parent / "_internal" / "assets" / "dashboard",
                ])
            frame_files = []
            for root in roots:
                if not root.exists():
                    continue
                # f0001.jpg style from ffmpeg
                found = sorted(root.glob("f*.jpg"))
                if not found:
                    found = sorted(root.glob("bg*.jpg"), key=lambda p: p.name)
                if found:
                    frame_files = found
                    break
            for fp in frame_files:
                try:
                    self._dash_photos.append(Image.open(fp).convert("RGB"))
                except Exception:
                    pass
            self._dash_idx = 0
            self._dash_size = (0, 0)
            self._dash_scaled = []  # cache scaled frames for current canvas size

        self.update_idletasks()
        cw = max(int(self.canvas.winfo_width()), 400)
        ch = max(int(self.canvas.winfo_height()), 300)
        if cw < 50:
            cw, ch = 900, 600

        self._draw_dashboard_frame(cw, ch)
        self._schedule_dashboard_cycle()

    def _ensure_scaled_frames(self, cw, ch):
        """Cover-scale all frames to canvas size (cached until resize)."""
        if self._dash_scaled and getattr(self, "_dash_size", None) == (cw, ch):
            return
        self._dash_size = (cw, ch)
        self._dash_scaled = []
        if not self._dash_photos:
            return
        for im0 in self._dash_photos:
            im = im0.copy()
            iw, ih = im.size
            scale = max(cw / max(iw, 1), ch / max(ih, 1))
            nw, nh = max(1, int(iw * scale)), max(1, int(ih * scale))
            im = im.resize((nw, nh), Image.Resampling.LANCZOS)
            left = max(0, (nw - cw) // 2)
            top = max(0, (nh - ch) // 2)
            im = im.crop((left, top, left + cw, top + ch))
            self._dash_scaled.append(im)

    def _draw_dashboard_frame(self, cw=None, ch=None):
        if cw is None or ch is None:
            cw = max(int(self.canvas.winfo_width()), 400)
            ch = max(int(self.canvas.winfo_height()), 300)
        self.canvas.delete("dash")
        if self._dash_photos:
            self._ensure_scaled_frames(cw, ch)
            if self._dash_scaled:
                im = self._dash_scaled[self._dash_idx % len(self._dash_scaled)]
                self._dash_bg_photo = ImageTk.PhotoImage(im)
                self.canvas.create_image(0, 0, anchor=tk.NW, image=self._dash_bg_photo, tags="dash")
                self.canvas.config(scrollregion=(0, 0, cw, ch), bg="#0b1220")
        else:
            self.canvas.config(bg="#1a2744")
            self.canvas.create_rectangle(0, 0, cw, ch, fill="#1a2744", outline="", tags="dash")

        # Transparent text only
        cx, cy = cw // 2, ch // 2
        line1, line2 = "Drag & Drop", "your file here"
        for dx, dy in ((1, 1),):
            self.canvas.create_text(cx + dx, cy - 10 + dy, text=line1, fill="#000000",
                                    font=("Segoe UI", 15, "bold"), tags="dash")
            self.canvas.create_text(cx + dx, cy + 12 + dy, text=line2, fill="#000000",
                                    font=("Segoe UI", 11), tags="dash")
        self.canvas.create_text(cx, cy - 10, text=line1, fill="#1a1a1a",
                                font=("Segoe UI", 15, "bold"), tags="dash")
        self.canvas.create_text(cx, cy + 12, text=line2, fill="#1a1a1a",
                                font=("Segoe UI", 11), tags="dash")

    def _schedule_dashboard_cycle(self):
        self._stop_dashboard(cancel_only=True)
        # ~12 fps to match source video
        if not self.pdf.doc and self._dash_photos:
            self._dash_job = self.after(125, self._cycle_dashboard)

    def _cycle_dashboard(self):
        self._dash_job = None
        if self.pdf.doc or not self._dash_photos:
            return
        try:
            self._dash_idx = (self._dash_idx + 1) % len(self._dash_photos)
            self._draw_dashboard_frame()
        except Exception:
            return
        if not self.pdf.doc:
            self._schedule_dashboard_cycle()

    def _stop_dashboard(self, cancel_only=False):
        if self._dash_job is not None:
            try:
                self.after_cancel(self._dash_job)
            except Exception:
                pass
            self._dash_job = None

    def _on_drop_files(self, event):
        """Handle drag-and-drop of files onto the window/canvas."""
        try:
            data = event.data
        except Exception:
            data = str(event)
        paths = []
        if isinstance(data, str):
            raw = data.strip()
            if raw.startswith("{"):
                import re as _re
                paths = _re.findall(r"\{([^}]+)\}", raw)
                if not paths:
                    paths = [raw.strip("{}")]
            else:
                paths = [raw]
        for p in paths:
            p = p.strip().strip("{}")
            if os.path.isfile(p):
                self.after(0, lambda path=p: self._safe_load_dropped(path))
                break

    def _decode_drop_path(self, f):
        try:
            if isinstance(f, bytes):
                s = None
                for enc in ("utf-8", "mbcs", "cp1252", "latin-1"):
                    try:
                        s = f.decode(enc)
                        break
                    except Exception:
                        pass
                if not s:
                    try:
                        s = os.fsdecode(f)
                    except Exception:
                        s = f.decode("utf-8", errors="ignore")
            else:
                s = str(f)
            s = s.replace("\x00", "").strip().strip('"').strip("'")
            # also strip real nulls
            s = "".join(ch for ch in s if ch != "\0" and ord(ch) != 0)
            if s.startswith("{") and s.endswith("}"):
                s = s[1:-1]
            return s
        except Exception:
            return ""

    def _on_windnd_drop(self, files):
        """Windnd runs off the Tk thread — only enqueue paths, never touch Tk here."""
        try:
            if files is None:
                return
            raw = files
            if isinstance(raw, (str, bytes)):
                raw = [raw]
            try:
                iterable = list(raw)
            except Exception:
                return
            paths = []
            for f in iterable:
                p = self._decode_drop_path(f)
                if p and os.path.isfile(p):
                    paths.append(p)
            if paths:
                self._drop_queue.put(paths)
        except Exception:
            pass

    def _poll_drop_queue(self):
        """UI-thread poller for dropped files (prevents windnd/Tk cross-thread crash)."""
        try:
            if not getattr(self, "_drop_busy", False):
                paths = None
                try:
                    paths = self._drop_queue.get_nowait()
                except Exception:
                    paths = None
                if paths:
                    self._drop_busy = True
                    try:
                        self._safe_load_dropped_list(paths)
                    finally:
                        self._drop_busy = False
        except Exception:
            self._drop_busy = False
        try:
            self.after(250, self._poll_drop_queue)
        except Exception:
            pass

    def _safe_load_dropped_list(self, paths):
        try:
            if not paths:
                return
            try:
                self._stop_dashboard()
            except Exception:
                pass
            self.screenshot_mode = False
            self.fill_mode = False
            self.copy_sign_mode = False
            self.placing_signature = False
            self.placing_symbol = False
            self.link_mode = False
            self._move_active = False
            for path in list(paths)[: self.MAX_TABS]:
                try:
                    if not path or not os.path.isfile(path):
                        continue
                    self._load_path(path)
                except Exception as e:
                    try:
                        messagebox.showerror("Open failed", str(e), parent=self)
                    except Exception:
                        pass
                    break
        except Exception as e:
            try:
                messagebox.showerror("Drop failed", str(e), parent=self)
            except Exception:
                pass

    def _safe_load_dropped(self, path):
        self._safe_load_dropped_list([path] if path else [])

    def _on_canvas_configure(self, event):
        if self.pdf.doc:
            return
        if event.width > 50 and event.height > 50:
            self._dash_size = (0, 0)
            self._dash_scaled = []
            if self._dash_job is None:
                def _redraw():
                    if not self.pdf.doc:
                        self._draw_dashboard_frame(event.width, event.height)
                        self._schedule_dashboard_cycle()
                self._dash_job = self.after(150, _redraw)

    def _ask_text_with_bangla(self, prompt="Enter text:"):
        """Text entry dialog with optional on-screen Bangla keyboard."""
        win = tk.Toplevel(self)
        win.title("Type text")
        win.configure(bg=COLORS["surface"])
        win.transient(self)
        win.grab_set()
        result = {"value": None}
        tk.Label(win, text=prompt, bg=COLORS["surface"], fg=COLORS["text"],
                 font=("Segoe UI", 10)).pack(padx=12, pady=(12, 4))
        var = tk.StringVar()
        entry = tk.Entry(win, textvariable=var, width=40, font=("Segoe UI", 12),
                         bg=COLORS["surface2"], fg=COLORS["text"],
                         insertbackground=COLORS["text"], relief=tk.FLAT)
        entry.pack(padx=12, pady=6)
        entry.focus_set()

        def insert_char(ch):
            if ch == "BACK":
                cur = var.get()
                var.set(cur[:-1])
            else:
                entry.insert(tk.INSERT, ch)
            entry.focus_set()

        def ok():
            result["value"] = var.get()
            win.destroy()

        def cancel():
            win.destroy()

        bf = tk.Frame(win, bg=COLORS["surface"])
        bf.pack(pady=8)
        tk.Button(bf, text="বাংলা কিবোর্ড", command=lambda: BanglaKeyboard(win, on_char=insert_char, on_enter=ok),
                  bg="#0ea5e9", fg="white", relief=tk.FLAT, padx=10, pady=4,
                  font=("Segoe UI", 9), cursor="hand2").pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="OK", command=ok, bg=COLORS["accent"], fg="white",
                  relief=tk.FLAT, padx=14, pady=4, font=("Segoe UI", 9), cursor="hand2").pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="Cancel", command=cancel, bg=COLORS["surface2"], fg=COLORS["text"],
                  relief=tk.FLAT, padx=10, pady=4, font=("Segoe UI", 9), cursor="hand2").pack(side=tk.LEFT, padx=4)
        entry.bind("<Return>", lambda e: ok())
        win.wait_window()
        return result["value"]

    def _open_bangla_kb_standalone(self):
        """Bangla keyboard that fills a buffer — paste into edit fields."""
        buf = {"text": ""}

        def insert_char(ch):
            if ch == "BACK":
                buf["text"] = buf["text"][:-1]
            else:
                buf["text"] += ch
            self.status.config(text="বাংলা: " + buf["text"][-40:])

        def on_enter():
            # copy to clipboard for easy paste
            try:
                self.clipboard_clear()
                self.clipboard_append(buf["text"])
                messagebox.showinfo(
                    "Copied",
                    "বাংলা টেক্সট কপি হয়েছে।\n\nএখন টেক্সটে ডাবল-ক্লিক করে Edit → Ctrl+V পেস্ট করুন।\n\n" + buf["text"][:200],
                    parent=self,
                )
            except Exception as e:
                messagebox.showinfo("Text", buf["text"] or "(empty)", parent=self)

        BanglaKeyboard(self, on_char=insert_char, on_enter=on_enter)


    def open_pdf_merger(self):
        PDFMergerWindow(self)

    def print_document(self):
        """Open Windows print dialog for the current PDF."""
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a file first.", parent=self)
            return
        try:
            tmp = Path(tempfile.gettempdir()) / f"OnePDF_print_{os.getpid()}.pdf"
            self.pdf.doc.save(str(tmp), garbage=1, deflate=True)
            path = str(tmp)
            if not sys.platform.startswith("win"):
                messagebox.showinfo("Print", f"Saved print copy to:\n{path}", parent=self)
                return
            opened = False
            # 1) ShellExecute Print verb
            try:
                import ctypes
                rc = ctypes.windll.shell32.ShellExecuteW(
                    None, "print", path, None, None, 1
                )
                if rc > 32:
                    opened = True
            except Exception:
                pass
            # 2) PowerShell Start-Process -Verb Print
            if not opened:
                try:
                    import subprocess
                    subprocess.Popen(
                        [
                            "powershell",
                            "-NoProfile",
                            "-Command",
                            f'Start-Process -FilePath "{path}" -Verb Print',
                        ],
                        shell=False,
                    )
                    opened = True
                except Exception:
                    pass
            # 3) os.startfile print
            if not opened:
                try:
                    os.startfile(path, "print")
                    opened = True
                except Exception:
                    pass
            # 4) open file so user can print from viewer
            if not opened:
                try:
                    os.startfile(path)
                    opened = True
                except Exception as e:
                    messagebox.showerror("Print failed", str(e), parent=self)
                    return
            self.status.config(text="Print dialog requested — choose printer in Windows")
        except Exception as e:
            messagebox.showerror("Print failed", str(e), parent=self)

    def start_place_symbol(self, char):
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a file first.", parent=self)
            return
        self.placing_symbol = True
        self.symbol_char = char
        self.fill_mode = False
        self.screenshot_mode = False
        self.copy_sign_mode = False
        self.placing_signature = False
        self._move_active = False
        try:
            self.canvas.config(cursor="crosshair")
        except Exception:
            pass
        self.status.config(text=f"Symbol '{char}' — click on the page to place it")

    def _place_symbol_at(self, pt):
        char = self.symbol_char or "✓"
        self.placing_symbol = False
        self.symbol_char = None
        try:
            self.canvas.config(cursor="")
        except Exception:
            pass
        page = self.pdf.get_page(self.current_page)
        if not page:
            return
        self.pdf.save_state()
        try:
            # Prefer a Unicode-capable font if available (Bangla pack covers many symbols poorly;
            # use built-in or Noto if needed). Helvetica misses many symbols — try fontfile.
            fontfile = get_bengali_font_path()
            fontsize = 14
            origin = fitz.Point(pt.x, pt.y)
            if fontfile:
                try:
                    page.insert_font(fontname="symfont", fontfile=fontfile)
                    page.insert_text(origin, char, fontname="symfont", fontfile=fontfile,
                                     fontsize=fontsize, color=self.text_color, overlay=True)
                except Exception:
                    page.insert_text(origin, char, fontname="helv", fontsize=fontsize,
                                     color=self.text_color, overlay=True)
            else:
                page.insert_text(origin, char, fontname="helv", fontsize=fontsize,
                                 color=self.text_color, overlay=True)
            self.pdf.dirty = True
            self.render_page()
            self.status.config(text=f"Symbol placed: {char}")
        except Exception as e:
            messagebox.showerror("Symbol", str(e), parent=self)

    def _tk_exception_guard(self, exc_type, exc_value, exc_tb):
        try:
            self.status.config(text=f"Error: {exc_value}")
        except Exception:
            pass

    def show_about(self):
        messagebox.showinfo(
            "About",
            f"{APP_NAME}\nVersion {APP_VERSION}  •  {APP_YEAR}\n\n"
            f"Made by {APP_STUDIO}\n\n"
            "Offline PDF, Image & Document editor.\n"
            "View • Edit text in place • Sign • Screenshot\n"
            "Color Fill Box • Copy signature from PDF\n"
            "Open PDF / Images / Word → Save as PDF\n\n"
            "No internet required.",
        )

    def _on_close(self):
        if self.pdf.dirty:
            if not messagebox.askyesno("Quit", "Unsaved changes will be lost. Quit?"):
                return
        self.pdf.close()
        self.destroy()


def main():
    try:
        from ctypes import windll
        windll.shcore.SetProcessDpiAwareness(1)
    except Exception:
        pass
    app = OnePDFEditor()
    # Windows "Open with" / double-click association: path comes as argv
    if len(sys.argv) > 1:
        arg_path = sys.argv[1].strip('"')
        if os.path.isfile(arg_path):
            app.after(200, lambda p=arg_path: app._load_path(p))
    app.mainloop()


if __name__ == "__main__":
    main()
